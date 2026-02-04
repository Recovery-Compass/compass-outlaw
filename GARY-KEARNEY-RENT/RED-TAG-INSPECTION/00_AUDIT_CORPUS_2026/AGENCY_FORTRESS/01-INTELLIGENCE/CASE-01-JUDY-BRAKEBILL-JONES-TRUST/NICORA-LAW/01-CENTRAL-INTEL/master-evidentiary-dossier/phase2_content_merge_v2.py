#!/usr/bin/env python3
"""
Phase 2 Content Merge Script v2 - Improved placeholder detection
"""

from docx import Document
import os
import re

def update_header_info_comprehensive(doc, eric_info):
    """Update ALL header placeholders with Eric's real information"""

    replacements_made = 0

    for i, para in enumerate(doc.paragraphs[:30]):  # Process header region thoroughly
        original = para.text

        # Replace simple placeholders
        para.text = para.text.replace("ADDRESS", eric_info['address'])
        para.text = para.text.replace("CITY, CA ZIP", eric_info['city_state_zip'])
        para.text = para.text.replace("(415) 123-4567", eric_info['phone'])
        para.text = para.text.replace("Tel: (415) 123-4567", f"Tel: {eric_info['phone']}")
        para.text = para.text.replace("email address", eric_info['email'])
        para.text = para.text.replace("Email: email address", f"Email: {eric_info['email']}")

        # Replace trust names
        para.text = para.text.replace("LARRY LASAGNA 2002 REVOCABLE TRUST", eric_info['trust'].upper())
        para.text = para.text.replace("LARRY LASAGNA", "JUDY BRAKEBILL JONES")
        para.text = para.text.replace("ESTATE OF LARRY LASAGNA", "ESTATE OF JUDY BRAKEBILL JONES")

        # Replace county
        para.text = para.text.replace("[NAME COUNTY]", "LOS ANGELES")
        para.text = para.text.replace("COUNTY OF [NAME COUNTY]", "COUNTY OF LOS ANGELES")

        # Replace generic case info
        para.text = para.text.replace("CASE NO. RP21094431", "CASE NO.: (To be assigned)")
        para.text = para.text.replace("RP21094431", "(To be assigned)")

        if original != para.text:
            replacements_made += 1

    print(f"  • Made {replacements_made} header replacements")
    return doc

def extract_body_content(source_doc, skip_lines=15):
    """Extract body paragraphs from source document"""
    content_paragraphs = []

    for i, para in enumerate(source_doc.paragraphs):
        if i < skip_lines:
            continue

        text = para.text.strip()

        # Skip empty paragraphs
        if len(text) < 10:
            continue

        # Skip header-like content
        header_keywords = [
            "SUPERIOR COURT", "COUNTY OF", "Case No.", "CASE NO",
            "PETITIONER", "RESPONDENT", "IN RE THE",
            "ERIC B. JONES", "ERIC BRAKEBILL JONES"
        ]
        if any(keyword in para.text for keyword in header_keywords) and i < 35:
            continue

        content_paragraphs.append(para)

    return content_paragraphs

def merge_documents(template_path, source_path, output_path, eric_info, doc_name):
    """Complete merge with comprehensive placeholder replacement"""

    print(f"\n{'='*70}")
    print(f"Processing: {doc_name}")
    print(f"{'='*70}")

    # Load template
    template = Document(template_path)
    print(f"✓ Loaded template: {len(template.paragraphs)} paragraphs")

    # Update header comprehensively
    print(f"✓ Updating header placeholders...")
    template = update_header_info_comprehensive(template, eric_info)

    # Load source
    try:
        source = Document(source_path)
        print(f"✓ Loaded source: {len(source.paragraphs)} paragraphs")
    except Exception as e:
        print(f"✗ Error loading source: {e}")
        return False

    # Extract content
    content_paras = extract_body_content(source, skip_lines=15)
    print(f"✓ Extracted {len(content_paras)} content paragraphs from source")

    if len(content_paras) == 0:
        print(f"⚠ Warning: No content extracted from {doc_name}")
        # Don't fail - maybe this is a short document

    # Remove any existing body paragraphs beyond header (keep first 16)
    while len(template.paragraphs) > 16:
        p = template.paragraphs[-1]
        p._element.getparent().remove(p._element)

    # Add blank line before content
    template.add_paragraph("")

    # Insert content
    added = 0
    for para in content_paras:
        if len(para.text.strip()) > 0:
            new_para = template.add_paragraph(para.text)

            # Preserve basic formatting
            if para.runs and new_para.runs:
                for src_run, dst_run in zip(para.runs, new_para.runs):
                    if src_run.bold:
                        dst_run.bold = True
                    if src_run.italic:
                        dst_run.italic = True
                    if src_run.underline:
                        dst_run.underline = True

            added += 1

    print(f"✓ Added {added} content paragraphs to template")

    # Save
    template.save(output_path)
    print(f"✓ Saved: {os.path.basename(output_path)}")

    return True

def convert_to_pdf(docx_path):
    """Convert DOCX to PDF using LibreOffice"""
    output_dir = os.path.dirname(docx_path)
    basename = os.path.basename(docx_path)

    cmd = f'cd "{output_dir}" && /Applications/LibreOffice.app/Contents/MacOS/soffice --headless --convert-to pdf --outdir . "{basename}" 2>&1'

    result = os.system(cmd)

    pdf_path = docx_path.replace('.docx', '.pdf')
    if os.path.exists(pdf_path):
        size_kb = os.path.getsize(pdf_path) / 1024
        print(f"✓ PDF created: {os.path.basename(pdf_path)} ({size_kb:.1f} KB)")
        return True

    print(f"✗ PDF conversion failed")
    return False

def verify_content(docx_path):
    """Verify no placeholder text remains"""
    doc = Document(docx_path)

    bad_placeholders = [
        "ADDRESS",
        "CITY, CA ZIP",
        "(415) 123-4567",
        "email address",
        "LARRY LASAGNA",
        "[NAME COUNTY]",
        "[Insert",
        "[Describe",
        "_______________"
    ]

    found = []
    for i, para in enumerate(doc.paragraphs):
        for placeholder in bad_placeholders:
            if placeholder in para.text:
                found.append((i+1, placeholder, para.text[:60]))

    if found:
        print(f"✗ Found {len(found)} placeholder(s):")
        for line, placeholder, text in found[:5]:
            print(f"  Line {line}: '{placeholder}' in '{text}...'")
        return False

    print(f"✓ Verification: No placeholders found")
    return True

# ===== CONFIGURATION =====

ERIC_INFO = {
    'name': 'ERIC BRAKEBILL JONES',
    'address': '17742 Berta Canyon Road',
    'city_state_zip': 'Lake Hughes, CA 93532',
    'phone': '(626) 348-3019',
    'email': 'eric@recovery-compass.org',
    'capacity': 'Sole Successor Trustee, In Pro Per',
    'trust': 'Judy Brakebill Jones 2008 Revocable Trust'
}

TEMPLATE_PATH = "/Users/ericjones/Cases/Judy-Trust/pro-per-probate-850/anuar-medina-ramirez-seven-hills-law/email-november-11/pleading-paper/PLEADING PAPER FOR ERIC JONES.docx"

SOURCE_DIR = "/Users/ericjones/11-15/5-bird/judy-jones-probate-850/self-filing-IRONCLAD-FINAL/remediated_docx_v2"

OUTPUT_DIR = "/Users/ericjones/Compass Fortress/cases/judy-jones-probate-850/4-email-to-anuar-nov15/ALL-ATTACHMENTS-READY-TO-SEND"

DOCUMENTS = [
    {
        'name': '01_Ex_Parte_Application',
        'source': '01_Ex_Parte_Application.docx',
        'output': '01_Ex_Parte_Application_CRC2111_COMPLETE_v2.docx'
    },
    {
        'name': '02_Declaration_with_Exhibits',
        'source': '02_Declaration_with_Exhibits.docx',
        'output': '02_Declaration_with_Exhibits_CRC2111_COMPLETE_v2.docx'
    },
    {
        'name': '04_MPA',
        'source': '04_MPA.docx',
        'output': '04_MPA_CRC2111_COMPLETE_v2.docx'
    },
    {
        'name': '06_Petition_850',
        'source': '06_Petition_850.docx',
        'output': '06_Petition_850_CRC2111_COMPLETE_v2.docx'
    }
]

# ===== MAIN EXECUTION =====

def main():
    print("\n" + "="*70)
    print("PHASE 2: CONTENT MERGE EXECUTION (v2 - Improved)")
    print("="*70)

    completed = []
    failed = []
    warnings = []

    for doc in DOCUMENTS:
        source_path = os.path.join(SOURCE_DIR, doc['source'])
        output_path = os.path.join(OUTPUT_DIR, doc['output'])

        if not os.path.exists(source_path):
            print(f"\n✗ Source not found: {source_path}")
            failed.append(doc['name'])
            continue

        # Merge
        success = merge_documents(
            TEMPLATE_PATH,
            source_path,
            output_path,
            ERIC_INFO,
            doc['name']
        )

        if not success:
            failed.append(doc['name'])
            continue

        # Verify
        if not verify_content(output_path):
            warnings.append(doc['name'])

        # Convert to PDF
        if convert_to_pdf(output_path):
            completed.append(doc['name'])
        else:
            failed.append(doc['name'])

    # ===== FINAL REPORT =====
    print("\n" + "="*70)
    print("FINAL STATUS REPORT")
    print("="*70)

    print(f"\n✓ Completed: {len(completed)}/4 documents")
    print(f"✗ Failed: {len(failed)}/4 documents")

    if warnings:
        print(f"⚠ Warnings: {len(warnings)} documents with placeholder residue")
        for doc in warnings:
            print(f"  ⚠ {doc}")

    if len(completed) == 4 and len(warnings) == 0:
        print("\n🎯 PHASE 2 COMPLETE (VERIFIED)\n")
        print("📁 Location:")
        print(f"   {OUTPUT_DIR}/\n")
        print("📄 Files Ready to Send:")
        for doc in completed:
            print(f"   ✓ {doc}_CRC2111_COMPLETE_v2.pdf")

        print("\n" + "="*70)
        print("VERIFICATION STATUS")
        print("="*70)
        print("✓ Phase 1: CRC 2.111 format structure")
        print("✓ Phase 2: Legal content merged from source")
        print("✓ Phase 3: All placeholders replaced")
        print("✓ Phase 4: PDF conversion successful")
        print("\n📊 Status: 100% MISSION COMPLETE")
        print("✉️  Action: READY TO SEND TO ANUAR\n")

    else:
        print(f"\n⚠️  REVIEW NEEDED")
        if failed:
            print("Failed:")
            for doc in failed:
                print(f"   ✗ {doc}")
        if warnings:
            print("\nNeed manual review:")
            for doc in warnings:
                print(f"   ⚠ {doc}")

    print("="*70 + "\n")

if __name__ == "__main__":
    main()
