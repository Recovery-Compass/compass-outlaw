import os
from docx import Document

BASE_DIR = "/Users/ericjones/Fortress/02-CASES/Judy-Trust/GENERATED_FILINGS"
MASTER_TEMPLATE = os.path.join(BASE_DIR, "MASTER_TEMPLATE.docx")

# --- OPTIMIZED ARGUMENT (Exhibit C Correction) ---
NOTICE_TEXT = """I, ERIC BRAKEBILL JONES, declare:

1. I am the Petitioner and Successor Trustee in this matter.

2. GOOD CAUSE EXCEPTION TO NOTICE: I have not given notice to the Respondents or their counsel regarding this Ex Parte Application. Providing notice would defeat the purpose of this application and precipitate the immediate harm I seek to prevent.

3. As detailed in my Declaration in Support (Exhibit C), Respondents’ counsel has already admitted to the unauthorized sale of Trust assets ($10,650). If given notice of this hearing before the freezing order is in place, there is a high probability they will immediately dissipate the remaining cash assets or conceal the digital devices (iPhone) required to save the real property from foreclosure.

4. I request that the Court waive notice requirements pursuant to CRC 3.1204(b)(3) to prevent immediate and irreparable harm."""

def main():
    if not os.path.exists(MASTER_TEMPLATE):
        print("❌ Error: Master Template not found. Please ensure you saved 'MASTER_TEMPLATE.docx'.")
        return

    try:
        doc = Document(MASTER_TEMPLATE)
        
        # 1. Update Title
        NEW_TITLE = "DECLARATION OF ERIC B. JONES RE: GOOD CAUSE EXCEPTION TO NOTICE"
        
        # Replace Title in Body
        for p in doc.paragraphs:
            if "DECLARATION OF ERIC" in p.text:
                p.text = NEW_TITLE
        
        # Replace Title in Footer
        for section in doc.sections:
            for f in section.footer.paragraphs:
                 if "DECLARATION" in f.text:
                     f.text = "DECLARATION RE NOTICE"

        # 2. Replace Body Text
        # We overwrite the first substantive paragraph found
        found_body = False
        for p in doc.paragraphs:
            if "1. Competence" in p.text or "4. Deposit" in p.text:
                p.text = NOTICE_TEXT
                found_body = True
                break
        
        if not found_body:
            doc.add_paragraph(NOTICE_TEXT)

        # Save
        doc.save(os.path.join(BASE_DIR, "03_Declaration_Re_Notice.docx"))
        print("✅ Created: 03_Declaration_Re_Notice.docx")

    except Exception as e:
        print(f"❌ Error: {e}")

if __name__ == "__main__":
    main()
