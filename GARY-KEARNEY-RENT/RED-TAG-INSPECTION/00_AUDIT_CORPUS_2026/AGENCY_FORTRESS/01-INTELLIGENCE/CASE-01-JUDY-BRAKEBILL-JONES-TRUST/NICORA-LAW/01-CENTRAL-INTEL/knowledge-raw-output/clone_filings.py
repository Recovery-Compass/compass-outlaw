import os
from docx import Document

# --- PATHS ---
BASE_DIR = "/Users/ericjones/Fortress/02-CASES/Judy-Trust/GENERATED_FILINGS"
MASTER_TEMPLATE = os.path.join(BASE_DIR, "MASTER_TEMPLATE.docx")

# --- IRONCLAD CONTENT ---
# This is the specific legal text for the other documents
MPA_TEXT = """II. LEGAL ARGUMENT

A. Unauthorized Sale and False Claim of Authority Under Probate Code § 859
Probate Code § 859 authorizes a court to award double damages when property has been taken or concealed from an estate in bad faith. Here, Respondent Heidi Blanchard and her counsel, Nicora Law Offices, LLP, have made dispositive admissions. (See Exhibit C).

B. Concealment of the Decedent's iPhone Violates RUFADAA
Respondents have concealed the decedent's original iPhone and provided a non-functional "decoy phone". This violates Probate Code §§ 870-872 and blocks access to Trust assets."""

TRO_TEXT = """TO RESPONDENTS HEIDI MARICHEN JONES BLANCHARD AND GARY WILLIAM JONES:

YOU ARE HEREBY ORDERED TO APPEAR before this Court to show cause why a preliminary injunction should not be issued.

PENDING THE HEARING, IT IS ORDERED THAT:
1. Respondents are restrained from spending, transferring, or encumbering the $10,650 proceeds from the sale of the 2003 Mercedes-Benz.
2. Respondents shall immediately deliver the decedent's authentic iPhone to Petitioner."""

PETITION_TEXT = """Petitioner ERIC BRAKEBILL JONES alleges:
1. Standing. Petitioner is the Sole Successor Trustee of the Judy Brakebill Jones 2008 Revocable Trust.
2. The Property. The Trust owns the real property at 17742 Berta Canyon Road, Salinas, CA.
3. The Theft. Respondents sold the Trust's Mercedes-Benz without authority.
4. Relief. Petitioner requests an order under Probate Code § 850 confirming Trust assets and § 859 double damages."""

# --- DOCUMENTS TO GENERATE ---
DOCS = [
    {"name": "04_MPA.docx", "title": "MEMORANDUM OF POINTS AND AUTHORITIES", "body": MPA_TEXT},
    {"name": "05_Proposed_TRO_OSC.docx", "title": "[PROPOSED] TEMPORARY RESTRAINING ORDER", "body": TRO_TEXT},
    {"name": "06_Petition_850.docx", "title": "PETITION FOR ORDER UNDER PROBATE CODE § 850", "body": PETITION_TEXT}
]

def main():
    if not os.path.exists(MASTER_TEMPLATE):
        print(f"❌ Error: Could not find {MASTER_TEMPLATE}")
        return

    print(f"Using Master Template: {MASTER_TEMPLATE}")

    for doc_info in DOCS:
        print(f"Cloning to create {doc_info['name']}...")
        try:
            doc = Document(MASTER_TEMPLATE)
            
            # 1. Update the Title (Header/Footer/Body)
            # We search for "DECLARATION" and replace it with the new title
            for p in doc.paragraphs:
                if "DECLARATION OF ERIC" in p.text:
                    p.text = doc_info['title']
            
            for section in doc.sections:
                for f in section.footer.paragraphs:
                    if "DECLARATION" in f.text:
                        f.text = doc_info['title']

            # 2. Replace the Body Text
            # We look for the start of the Declaration text ("1. Competence" or "4. Deposit")
            # And replace it with the new content
            found_body = False
            for p in doc.paragraphs:
                if "1. Competence" in p.text or "4. Deposit" in p.text:
                    p.text = doc_info['body']
                    found_body = True
                    # We stop looking after finding the start, but in a real clone 
                    # we might want to delete subsequent paragraphs. 
                    # For safety, we just overwrite the first major block.
                    break
            
            if not found_body:
                # Fallback: Append if we couldn't find the insertion point
                doc.add_paragraph(doc_info['body'])

            doc.save(os.path.join(BASE_DIR, doc_info['name']))
            print(f"✅ Created: {doc_info['name']}")

        except Exception as e:
            print(f"❌ Error: {e}")

if __name__ == "__main__":
    main()
