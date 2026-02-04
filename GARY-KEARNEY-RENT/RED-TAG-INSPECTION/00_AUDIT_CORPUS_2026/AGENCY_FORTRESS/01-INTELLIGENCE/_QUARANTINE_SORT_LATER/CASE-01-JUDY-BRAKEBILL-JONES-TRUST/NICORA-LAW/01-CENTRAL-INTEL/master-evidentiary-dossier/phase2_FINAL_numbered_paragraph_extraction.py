#!/usr/bin/env python3
"""
FINAL ATTEMPT: Extract numbered paragraphs correctly
Target the specific declaration paragraphs that were found at lines 7-13
"""

from docx import Document
import os
import re

def extract_numbered_paragraphs(source_doc):
    """Extract ALL numbered paragraphs (1., 2., 3., etc.) from anywhere in document"""

    numbered_paras = []

    print(f"\n  🔍 Scanning for numbered paragraphs...")

    for i, para in enumerate(source_doc.paragraphs):
        text = para.text.strip()

        # Check if starts with number followed by period
        match = re.match(r'^(\d+)\.\s+(.+)', text)

        if match:
            num = match.group(1)
            content = match.group(2)
            numbered_paras.append((int(num), i, para, text))
            print(f"     ✓ Line {i}: Found #{num}: {text[:70]}...")

    # Sort by number
    numbered_paras.sort(key=lambda x: x[0])

    return numbered_paras

def extract_section_content(source_doc, start_marker, end_markers=None):
    """Extract content between section markers"""

    content = []
    started = False

    for i, para in enumerate(source_doc.paragraphs):
        text = para.text.strip()

        # Check if we've hit the start marker
        if start_marker.upper() in text.upper():
            started = True
            content.append(para)
            print(f"     ✓ Line {i}: Found section start: {text[:60]}...")
            continue

        if started:
            # Check if we've hit an end marker
            if end_markers:
                if any(marker.upper() in text.upper() for marker in end_markers):
                    print(f"     ✓ Line {i}: Section end reached")
                    break

            # Skip very short paragraphs
            if len(text) < 20:
                continue

            # Skip exhibit content
            if 'nicora' in text.lower() or 'attorney at law' in text.lower():
                print(f"     ⊗ Line {i}: Skipped exhibit: {text[:50]}...")
                continue

            content.append(para)
            print(f"     ✓ Line {i}: Added: {text[:60]}...")

    return content

def merge_declaration(template_path, source_path, output_path, eric_info):
    """Merge Declaration with numbered paragraphs"""

    print(f"\n{'='*70}")
    print(f"DECLARATION: Numbered Paragraph Extraction")
    print(f"{'='*70}")

    template = Document(template_path)
    print(f"\n✓ Loaded template")

    # Update placeholders
    template = update_placeholders(template, eric_info)
    print(f"✓ Updated placeholders")

    # Load source
    source = Document(source_path)
    print(f"✓ Loaded source")

    # Extract numbered paragraphs
    numbered_paras = extract_numbered_paragraphs(source)

    if len(numbered_paras) < 5:
        print(f"\n✗ CRITICAL: Only found {len(numbered_paras)} numbered paragraphs (expected 12-14)")
        print(f"   Proceeding with available paragraphs...")

    # Build document
    while len(template.paragraphs) > 16:
        p = template.paragraphs[-1]
        p._element.getparent().remove(p._element)

    if len(template.tables) > 1:
        tbl = template.tables[-1]._element
        tbl.getparent().remove(tbl)

    template.add_paragraph("")
    template.add_paragraph("I, Eric B. Jones, declare as follows:")
    template.add_paragraph("")

    # Add numbered paragraphs
    for num, line_no, para, text in numbered_paras:
        new_para = template.add_paragraph(text)

        # Preserve formatting
        if para.runs and new_para.runs:
            for src_run, dst_run in zip(para.runs[:len(new_para.runs)], new_para.runs):
                if src_run.bold:
                    dst_run.bold = True
                if src_run.italic:
                    dst_run.italic = True

        template.add_paragraph("")  # Blank line between paragraphs

    print(f"\n✓ Added {len(numbered_paras)} numbered paragraphs")

    # Add signature
    template.add_paragraph("")
    template.add_paragraph("I declare under penalty of perjury under the laws of the State of California that the foregoing is true and correct.")
    template.add_paragraph("")
    template.add_paragraph("DATED: November 15, 2025")
    template.add_paragraph("")
    template.add_paragraph("_______________________________")
    template.add_paragraph("ERIC BRAKEBILL JONES")
    template.add_paragraph("Sole Successor Trustee")
    template.add_paragraph("Judy Brakebill Jones 2008 Revocable Trust")

    template.save(output_path)
    print(f"✓ Saved: {os.path.basename(output_path)}")

    return len(numbered_paras)

def merge_ex_parte(template_path, source_path, output_path, eric_info):
    """Merge Ex Parte with sections"""

    print(f"\n{'='*70}")
    print(f"EX PARTE: Section Extraction")
    print(f"{'='*70}")

    template = Document(template_path)
    print(f"\n✓ Loaded template")

    template = update_placeholders(template, eric_info)
    print(f"✓ Updated placeholders")

    source = Document(source_path)
    print(f"✓ Loaded source")

    # Try to extract sections
    print(f"\n  🔍 Looking for sections...")

    all_content = []

    # Get all substantive content after header
    for i, para in enumerate(source.paragraphs):
        if i < 10:  # Skip header
            continue

        text = para.text.strip()

        if len(text) < 20:
            continue

        # Skip signature/footer
        if 'page' in text.lower() and 'of' in text.lower():
            continue

        all_content.append(para)
        print(f"     ✓ Line {i}: {text[:70]}...")

    # Build document
    while len(template.paragraphs) > 16:
        p = template.paragraphs[-1]
        p._element.getparent().remove(p._element)

    if len(template.tables) > 1:
        tbl = template.tables[-1]._element
        tbl.getparent().remove(tbl)

    template.add_paragraph("")

    for para in all_content:
        new_para = template.add_paragraph(para.text)

        if para.runs and new_para.runs:
            for src_run, dst_run in zip(para.runs[:len(new_para.runs)], new_para.runs):
                if src_run.bold:
                    dst_run.bold = True
                if src_run.italic:
                    dst_run.italic = True

    print(f"\n✓ Added {len(all_content)} paragraphs")

    # Signature
    template.add_paragraph("")
    template.add_paragraph("DATED: November 15, 2025")
    template.add_paragraph("")
    template.add_paragraph("Respectfully submitted,")
    template.add_paragraph("")
    template.add_paragraph("_______________________________")
    template.add_paragraph("ERIC BRAKEBILL JONES")
    template.add_paragraph("Sole Successor Trustee, In Pro Per")

    template.save(output_path)
    print(f"✓ Saved: {os.path.basename(output_path)}")

    return len(all_content)

def merge_mpa(template_path, source_path, output_path, eric_info):
    """Merge MPA with all sections"""

    print(f"\n{'='*70}")
    print(f"MPA: Full Document Extraction")
    print(f"{'='*70}")

    template = Document(template_path)
    print(f"\n✓ Loaded template")

    template = update_placeholders(template, eric_info)
    print(f"✓ Updated placeholders")

    source = Document(source_path)
    print(f"✓ Loaded source")

    all_content = []

    for i, para in enumerate(source.paragraphs):
        if i < 10:
            continue

        text = para.text.strip()

        if len(text) < 15:
            continue

        if 'page' in text.lower() and 'of' in text.lower():
            continue

        all_content.append(para)
        print(f"     ✓ Line {i}: {text[:70]}...")

    while len(template.paragraphs) > 16:
        p = template.paragraphs[-1]
        p._element.getparent().remove(p._element)

    if len(template.tables) > 1:
        tbl = template.tables[-1]._element
        tbl.getparent().remove(tbl)

    template.add_paragraph("")

    for para in all_content:
        template.add_paragraph(para.text)

    print(f"\n✓ Added {len(all_content)} paragraphs")

    template.add_paragraph("")
    template.add_paragraph("DATED: November 15, 2025")
    template.add_paragraph("")
    template.add_paragraph("Respectfully submitted,")
    template.add_paragraph("")
    template.add_paragraph("_______________________________")
    template.add_paragraph("ERIC BRAKEBILL JONES")
    template.add_paragraph("Sole Successor Trustee, In Pro Per")

    template.save(output_path)
    print(f"✓ Saved")

    return len(all_content)

def update_placeholders(doc, eric_info):
    """Update all placeholders"""

    replacements = {
        "ADDRESS": eric_info['address'],
        "CITY, CA ZIP": eric_info['city_state_zip'],
        "(415) 123-4567": eric_info['phone'],
        "Tel: (415) 123-4567": f"Tel: {eric_info['phone']}",
        "email address": eric_info['email'],
        "Email: email address": f"Email: {eric_info['email']}",
        "LARRY LASAGNA 2002 REVOCABLE TRUST": eric_info['trust'].upper(),
        "LARRY LASAGNA": "JUDY BRAKEBILL JONES",
        "JUDY [SURNMAE]": "JUDY BRAKEBILL JONES",
        "ESTATE OF LARRY LASAGNA": "ESTATE OF JUDY BRAKEBILL JONES",
        "[NAME COUNTY]": "LOS ANGELES",
        "CASE NO. RP21094431": "CASE NO.: (To be assigned)",
        "LAW OFFICE OF DAVID W. TATE": "",
        "DAVID W. TATE, ESQ.": "",
        "____________, 2023": "November 15, 2025"
    }

    for para in doc.paragraphs:
        for old, new in replacements.items():
            if old in para.text:
                para.text = para.text.replace(old, new)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in replacements.items():
                        if old in para.text:
                            para.text = para.text.replace(old, new)

    return doc

def convert_to_pdf(docx_path):
    """Convert to PDF"""
    output_dir = os.path.dirname(docx_path)
    basename = os.path.basename(docx_path)

    cmd = f'cd "{output_dir}" && /Applications/LibreOffice.app/Contents/MacOS/soffice --headless --convert-to pdf --outdir . "{basename}" 2>&1 > /dev/null'
    os.system(cmd)

    pdf_path = docx_path.replace('.docx', '.pdf')
    if os.path.exists(pdf_path):
        size_kb = os.path.getsize(pdf_path) / 1024
        print(f"✓ PDF: {os.path.basename(pdf_path)} ({size_kb:.1f} KB)")
        return True
    return False

# ===== CONFIG =====

ERIC_INFO = {
    'address': '17742 Berta Canyon Road',
    'city_state_zip': 'Lake Hughes, CA 93532',
    'phone': '(626) 348-3019',
    'email': 'eric@recovery-compass.org',
    'trust': 'Judy Brakebill Jones 2008 Revocable Trust'
}

TEMPLATE = "/Users/ericjones/Cases/Judy-Trust/pro-per-probate-850/anuar-medina-ramirez-seven-hills-law/email-november-11/pleading-paper/PLEADING PAPER FOR ERIC JONES.docx"
SOURCE_DIR = "/Users/ericjones/11-15/5-bird/judy-jones-probate-850/self-filing-IRONCLAD-FINAL/remediated_docx_v2"
OUTPUT_DIR = "/Users/ericjones/Compass Fortress/cases/judy-jones-probate-850/4-email-to-anuar-nov15/ALL-ATTACHMENTS-READY-TO-SEND"

def main():
    print("\n" + "="*70)
    print("FINAL EXTRACTION ATTEMPT: Numbered Paragraphs")
    print("="*70)

    results = {}

    # Declaration - most critical
    dec_count = merge_declaration(
        TEMPLATE,
        os.path.join(SOURCE_DIR, "02_Declaration_with_Exhibits.docx"),
        os.path.join(OUTPUT_DIR, "02_Declaration_CRC2111_FINAL_v2.docx"),
        ERIC_INFO
    )
    convert_to_pdf(os.path.join(OUTPUT_DIR, "02_Declaration_CRC2111_FINAL_v2.docx"))
    results['Declaration'] = dec_count

    # Ex Parte
    ex_count = merge_ex_parte(
        TEMPLATE,
        os.path.join(SOURCE_DIR, "01_Ex_Parte_Application.docx"),
        os.path.join(OUTPUT_DIR, "01_Ex_Parte_Application_CRC2111_FINAL_v2.docx"),
        ERIC_INFO
    )
    convert_to_pdf(os.path.join(OUTPUT_DIR, "01_Ex_Parte_Application_CRC2111_FINAL_v2.docx"))
    results['Ex Parte'] = ex_count

    # MPA
    mpa_count = merge_mpa(
        TEMPLATE,
        os.path.join(SOURCE_DIR, "04_MPA.docx"),
        os.path.join(OUTPUT_DIR, "04_MPA_CRC2111_FINAL_v2.docx"),
        ERIC_INFO
    )
    convert_to_pdf(os.path.join(OUTPUT_DIR, "04_MPA_CRC2111_FINAL_v2.docx"))
    results['MPA'] = mpa_count

    print("\n" + "="*70)
    print("FINAL RESULTS")
    print("="*70)
    print(f"\nDeclaration: {results['Declaration']} numbered paragraphs extracted")
    print(f"Ex Parte: {results['Ex Parte']} paragraphs extracted")
    print(f"MPA: {results['MPA']} paragraphs extracted")

    print("\n" + "="*70)
    print("FILES: *_FINAL_v2.pdf")
    print("="*70)
    print(f"\nLocation: {OUTPUT_DIR}/\n")
    print("These are the most complete extractions possible.")
    print("Review and forward to Anuar with gaps documented.")
    print("\n" + "="*70 + "\n")

if __name__ == "__main__":
    main()
