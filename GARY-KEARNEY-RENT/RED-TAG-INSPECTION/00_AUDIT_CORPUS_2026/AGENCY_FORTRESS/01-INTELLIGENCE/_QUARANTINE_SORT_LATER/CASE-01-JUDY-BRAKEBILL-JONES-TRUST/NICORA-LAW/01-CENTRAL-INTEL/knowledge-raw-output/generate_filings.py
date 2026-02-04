import os
from docx import Document
import datetime

# --- PATHS ---
TEMPLATE_PATH = 'pleading-paper/California-Pleading-Template-Google-Doc-Compatible.docx'
OUTPUT_DIR = 'GENERATED_FILINGS'

# --- DATA ---
TODAY = datetime.date.today().strftime("%B %d, 2025")
CONTACT_INFO = "ERIC BRAKEBILL JONES\nSuccessor Trustee, Pro Per\n5634 Noel Drive\nTemple City, CA 91780\n(626) 348-3019\neric@recovery-compass.org"

# --- IRONCLAD CONTENT ---
NICORA_TEXT = """4. Deposit of Proceeds and Dispositive Admission. On or about May 22, 2025, Respondents sold the Trust-owned Mercedes-Benz without my authorization as Successor Trustee. On May 24, 2025, Respondent Gretchen Angela Jones sent me an email acknowledging the unauthorized sale and attached a Chase Bank deposit slip showing $10,650 deposited into the personal account of Gary William Jones.

On June 27, 2025, Respondent Heidi Blanchard's attorney, Albert J. Nicora of Nicora Law Offices, LLP, sent me a formal letter that constitutes a dispositive admission of the unauthorized sale. The letter explicitly states that "Heidi, as Co-Executor, sold tangible personal property including a vehicle" and confirms the proceeds as $10,650.00. Critically, the letter falsely identified Heidi Blanchard as a "Co-Executor" (the Trust has none) and proposed unilateral distribution. This false claim of authority establishes bad faith under Probate Code § 859."""

FORECLOSURE_TEXT = """7. Immediate and Irreparable Harm. The Trust faces immediate and irreparable harm if a temporary restraining order is not issued. The risk is threefold: dissipation of assets, destruction of digital evidence, and foreclosure.

Critically, the concealed iPhone holds the keys to preventing the imminent foreclosure on the Trust's real property located at 17742 Berta Canyon Road, Salinas, CA. The reinstatement quote from Shellpoint Mortgage expires on December 3, 2025. Without the authentic iPhone to generate 2FA codes for bank access, I cannot execute the reinstatement payment. As early as August 5, 2025, I warned counsel of this risk. Respondents' continued obstruction has brought us to within days of this irreversible deadline."""

DOCS = [
    {"name": "01_Ex_Parte_Application.docx", "title": "EX PARTE APPLICATION FOR TEMPORARY RESTRAINING ORDER"},
    {"name": "02_Declaration_with_Exhibits.docx", "title": "DECLARATION OF ERIC BRAKEBILL JONES", "body": NICORA_TEXT + "\n\n" + FORECLOSURE_TEXT},
    {"name": "04_MPA.docx", "title": "MEMORANDUM OF POINTS AND AUTHORITIES"},
    {"name": "05_Proposed_TRO_OSC.docx", "title": "[PROPOSED] TEMPORARY RESTRAINING ORDER"},
    {"name": "06_Petition_850.docx", "title": "PETITION FOR ORDER UNDER PROBATE CODE § 850"}
]

def replace_text(doc, replacements):
    # 1. Paragraphs
    for p in doc.paragraphs:
        for key, val in replacements.items():
            if key in p.text:
                p.text = p.text.replace(key, val)
                # Center align venue/court if modified
                if "MONTEREY" in val:
                    p.alignment = 1

    # 2. Tables (Where Captions live)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, val in replacements.items():
                        if key in p.text:
                            p.text = p.text.replace(key, val)

    # 3. Headers/Footers
    for section in doc.sections:
        for h in section.header.paragraphs:
            for key, val in replacements.items():
                if key in h.text:
                    h.text = h.text.replace(key, val)
        for f in section.footer.paragraphs:
            for key, val in replacements.items():
                if key in f.text:
                    f.text = f.text.replace(key, val)

def main():
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)

    for doc_info in DOCS:
        print(f"Processing {doc_info['name']}...")
        doc = Document(TEMPLATE_PATH)
        
        # Define replacements for this specific doc
        replacements = {
            "SAN DIEGO": "MONTEREY",
            "NAME OF PLAINTIFF(S)": "ERIC BRAKEBILL JONES, Successor Trustee",
            "NAME OF DEFENDANT(S)": "HEIDI MARICHEN JONES BLANCHARD, et al.",
            "Plaintiff(s)": "Petitioner",
            "Defendant(s)": "Respondents",
            "Case No.:": "Case No.: [Assigned by Court]",
            "INSERT DOCUMENT TITLE": doc_info['title'],
            "DOCUMENT TITLE (e.g., NOTICE OF": doc_info['title'],
            "MOTION AND MOTION FOR STRIKING": "",
            "PORTIONS OF COMPLAINT)": "",
            "[Today’s Date]": TODAY,
            "YOUR NAME": "ERIC BRAKEBILL JONES",
            "Street Address": "", 
            "City, State, Zip": "",
            "Phone Number": "",
            "Fax Number": "",
            "Email:": "",
            "TYPE YOUR NAME": "Eric Brakebill Jones"
        }
        
        replace_text(doc, replacements)
        
        # Header Injection (Top Left)
        doc.paragraphs[0].text = CONTACT_INFO

        # Body Injection
        # Find where to put the text
        inserted = False
        for p in doc.paragraphs:
            if "Text of document" in p.text or "begins here" in p.text:
                p.text = doc_info['title'] + "\n\n" + doc_info.get('body', "")
                inserted = True
                break
        
        if not inserted:
            doc.add_paragraph(doc_info['title'])
            if 'body' in doc_info:
                doc.add_paragraph(doc_info['body'])

        doc.save(os.path.join(OUTPUT_DIR, doc_info['name']))
        print(f"✅ Fixed and Saved: {doc_info['name']}")

if __name__ == "__main__":
    main()
