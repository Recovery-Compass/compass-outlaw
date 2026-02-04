#!/usr/bin/env python3
"""
EMERGENCY FIX v2: Direct text replacement for Eric's address
"""

from docx import Document
import os

def fix_address_aggressive(doc):
    """Fix address by replacing text in all paragraphs"""

    fixes_made = []

    # Fix all paragraphs
    for i, para in enumerate(doc.paragraphs):
        original = para.text

        # Direct replacements
        if "17742 Berta Canyon Road" in para.text and ("Tel:" in para.text or "Email:" in para.text or i < 10):
            para.text = para.text.replace("17742 Berta Canyon Road", "5634 Noel Drive")
            fixes_made.append(f"Para {i}: 17742 Berta Canyon Road → 5634 Noel Drive")

        if "Lake Hughes, CA 93532" in para.text and ("Tel:" in para.text or "Email:" in para.text or i < 10):
            para.text = para.text.replace("Lake Hughes, CA 93532", "Temple City, CA 91780")
            fixes_made.append(f"Para {i}: Lake Hughes → Temple City")

        if "Lake Hughes" in para.text and "CA 93532" in para.text and i < 10:
            para.text = para.text.replace("Lake Hughes", "Temple City")
            para.text = para.text.replace("93532", "91780")
            fixes_made.append(f"Para {i}: Fixed Lake Hughes/93532")

    # Fix tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, para in enumerate(cell.paragraphs):
                    original = para.text

                    if "17742 Berta Canyon Road" in para.text:
                        para.text = para.text.replace("17742 Berta Canyon Road", "5634 Noel Drive")
                        fixes_made.append(f"Table {table_idx}, Row {row_idx}: 17742 Berta Canyon Road → 5634 Noel Drive")

                    if "Lake Hughes" in para.text and "93532" in para.text:
                        para.text = para.text.replace("Lake Hughes, CA 93532", "Temple City, CA 91780")
                        fixes_made.append(f"Table {table_idx}, Row {row_idx}: Lake Hughes → Temple City")

    return doc, fixes_made

def convert_to_pdf(docx_path):
    """Convert to PDF"""
    output_dir = os.path.dirname(docx_path)
    basename = os.path.basename(docx_path)

    cmd = f'cd "{output_dir}" && /Applications/LibreOffice.app/Contents/MacOS/soffice --headless --convert-to pdf --outdir . "{basename}" 2>&1 > /dev/null'
    os.system(cmd)

    pdf_path = docx_path.replace('.docx', '.pdf')
    return os.path.exists(pdf_path)

OUTPUT_DIR = "/Users/ericjones/Compass Fortress/cases/judy-jones-probate-850/4-email-to-anuar-nov15/ALL-ATTACHMENTS-READY-TO-SEND"

DOCUMENTS = [
    "01_Ex_Parte_Application_CRC2111_FINAL_v2.docx",
    "02_Declaration_CRC2111_FINAL_v2.docx",
    "04_MPA_CRC2111_FINAL_v2.docx"
]

def main():
    print("\n" + "="*70)
    print("EMERGENCY ADDRESS FIX v2 - AGGRESSIVE")
    print("="*70)
    print("\nCORRECT ADDRESS:")
    print("  Eric Brakebill Jones")
    print("  5634 Noel Drive")
    print("  Temple City, CA 91780")
    print("  (626) 348-3019")
    print("  eric@recovery-compass.org\n")

    all_successful = True

    for doc_file in DOCUMENTS:
        print(f"{'='*70}")
        print(f"Processing: {doc_file}")
        print(f"{'='*70}")

        docx_path = os.path.join(OUTPUT_DIR, doc_file)

        if not os.path.exists(docx_path):
            print(f"✗ File not found\n")
            all_successful = False
            continue

        # Load
        doc = Document(docx_path)
        print(f"✓ Loaded document")

        # Fix
        doc, fixes = fix_address_aggressive(doc)

        if fixes:
            print(f"✓ Made {len(fixes)} corrections:")
            for fix in fixes:
                print(f"  • {fix}")
        else:
            print(f"⚠ No corrections needed (address may already be correct)")

        # Save
        doc.save(docx_path)
        print(f"✓ Saved")

        # Convert
        if convert_to_pdf(docx_path):
            print(f"✓ PDF regenerated")
        else:
            print(f"✗ PDF conversion failed")
            all_successful = False

        print()

    print("="*70)
    if all_successful:
        print("✅ ALL DOCUMENTS FIXED")
    else:
        print("⚠️  SOME ISSUES OCCURRED")
    print("="*70)
    print("\nVerify PDFs now show:")
    print("  5634 Noel Drive")
    print("  Temple City, CA 91780")
    print("\n" + "="*70 + "\n")

if __name__ == "__main__":
    main()
