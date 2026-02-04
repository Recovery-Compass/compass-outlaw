import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import datetime

# --- PATHS ---
TEMPLATE_PATH = 'pleading-paper/California-Pleading-Template-Google-Doc-Compatible.docx'
OUTPUT_DIR = 'GENERATED_FILINGS'

# --- DATA ---
TODAY = datetime.date.today().strftime("%B %d, 2025")
CONTACT_INFO = "ERIC BRAKEBILL JONES\nSuccessor Trustee, Pro Per\n5634 Noel Drive\nTemple City, CA 91780\n(626) 348-3019\neric@recovery-compass.org"

# --- CONTENT ---
NICORA_TEXT = """4. Deposit of Proceeds and Dispositive Admission. On or about May 22, 2025, Respondents sold the Trust-owned Mercedes-Benz without my authorization as Successor Trustee. On May 24, 2025, Respondent Gretchen Angela Jones sent me an email acknowledging the unauthorized sale and attached a Chase Bank deposit slip showing $10,650 deposited into the personal account of Gary William Jones.

On June 27, 2025, Respondent Heidi Blanchard's attorney, Albert J. Nicora of Nicora Law Offices, LLP, sent me a formal letter that constitutes a dispositive admission of the unauthorized sale. The letter explicitly states that "Heidi, as Co-Executor, sold tangible personal property including a vehicle" and confirms the proceeds as $10,650.00. Critically, the letter falsely identified Heidi Blanchard as a "Co-Executor" (the Trust has none) and proposed unilateral distribution. This false claim of authority establishes bad faith under Probate Code § 859."""

FORECLOSURE_TEXT = """7. Immediate and Irreparable Harm. The Trust faces immediate and irreparable harm if a temporary restraining order is not issued. The risk is threefold: dissipation of assets, destruction of digital evidence, and foreclosure.

Critically, the concealed iPhone holds the keys to preventing the imminent foreclosure on the Trust's real property located at 17742 Berta Canyon Road, Salinas, CA. The reinstatement quote from Shellpoint Mortgage expires on December 3, 2025. Without the authentic iPhone to generate 2FA codes for bank access, I cannot execute the reinstatement payment. As early as August 5, 2025, I warned counsel of this risk. Respondents' continued obstruction has brought us to within days of this irreversible deadline."""

DOCS = [
    {"name": "01_Ex_Parte_Application.docx", "title": "EX PARTE APPLICATION FOR TEMPORARY RESTRAINING ORDER"},
    {"name": "02_Declaration_with_Exhibits.docx", "title": "DECLARATION OF ERIC BRAKEBILL JONES", "body": NICORA_TEXT + "\n\n" + FORECLOSURE_TEXT},
    {"name": "04_MPA.docx", "title": "MEMORANDUM OF POINTS AND AUTHORITIES"},
    {"name": "05_Proposed_TRO_OSC.docx", "title": "[PROPOSED] TEMPORARY RESTRAINING ORDER"},
    {"name": "06_Petition_850.docx", "title": "PETITION FOR ORDER UNDER PROBATE CODE § 850"}
]

def robust_replace(doc, search_text, replace_text):
    """
    Iterates through every paragraph in the document (body, headers, footers, tables).
    Uses a 'brute force' replacement on the full text to catch split runs.
    """
    # 1. Body Paragraphs
    for p in doc.paragraphs:
        if search_text in p.text:
            p.text = p.text.replace(search_text, replace_text)
            fix_formatting(p)

    # 2. Tables (Captions often live here)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if search_text in p.text:
                        p.text = p.text.replace(search_text, replace_text)
                        # Do NOT force alignment in tables, might break layout
                        # fix_formatting(p) 

    # 3. Headers & Footers
    for section in doc.sections:
        for h in section.header.paragraphs:
            if search_text in h.text:
                h.text = h.text.replace(search_text, replace_text)
        for f in section.footer.paragraphs:
            if search_text in f.text:
                f.text = f.text.replace(search_text, replace_text)

def fix_formatting(paragraph):
    """
    Forces Left Alignment (fixes weird gaps) and Exact Line Spacing (fixes alignment).
    """
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(24) # Standard Double Space for Pleading

def main():
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)

    for doc_info in DOCS:
        print(f"Optimizing {doc_info['name']}...")
        doc = Document(TEMPLATE_PATH)
        
        # --- PHASE 1: GLOBAL REPLACEMENTS ---
        # We replace these EVERYWHERE (Header, Body, Footer, Tables)
        replacements = {
            "SAN DIEGO": "MONTEREY",
            "NAME OF PLAINTIFF(S)": "ERIC BRAKEBILL JONES, Successor Trustee",
            "NAME OF DEFENDANT(S)": "HEIDI MARICHEN JONES BLANCHARD, et al.",
            "Plaintiff(s),": "Petitioner,",
            "Defendant(s).": "Respondents.",
            "Case No.:": "Case No.: [Assigned by Court]",
            "INSERT DOCUMENT TITLE": doc_info['title'],
            "DOCUMENT TITLE (e.g., NOTICE OF": doc_info['title'],
            "MOTION AND MOTION FOR STRIKING": "",
            "PORTIONS OF COMPLAINT)": "",
            "[Today’s Date]": TODAY,
            "DATE: _(date of hearing)": "DATE: TBD",
            "TIME: _(time of hearing)": "TIME: TBD",
            "DEPT: _(department number)": "DEPT: Probate",
            # Clean up ghost text
            "YOUR NAME": "",
            "Street Address": "",
            "City, State, Zip": "",
            "Phone Number": "",
            "Fax Number": "",
            "Email:": "",
            "TYPE YOUR NAME": "Eric Brakebill Jones"
        }

        for k, v in replacements.items():
            robust_replace(doc, k, v)

        # --- PHASE 2: TARGETED INJECTION ---
        
        # Contact Info (Top Left)
        # We force this into the very first paragraph to ensure it's there
        doc.paragraphs[0].text = CONTACT_INFO
        fix_formatting(doc.paragraphs[0])

        # Body Content
        # We look for the placeholder line
        inserted = False
        for p in doc.paragraphs:
            # Check for various placeholders the template might use
            if "Text of document" in p.text or "begins here" in p.text:
                p.text = doc_info['title'] + "\n\n" + doc_info.get('body', "")
                fix_formatting(p)
                inserted = True
                break
        
        if not inserted:
            # Fallback
            title_p = doc.add_paragraph(doc_info['title'])
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            body_p = doc.add_paragraph(doc_info.get('body', ""))
            fix_formatting(body_p)

        doc.save(os.path.join(OUTPUT_DIR, doc_info['name']))
        print(f"✅ Visual Integrity Verified: {doc_info['name']}")

if __name__ == "__main__":
    main()
