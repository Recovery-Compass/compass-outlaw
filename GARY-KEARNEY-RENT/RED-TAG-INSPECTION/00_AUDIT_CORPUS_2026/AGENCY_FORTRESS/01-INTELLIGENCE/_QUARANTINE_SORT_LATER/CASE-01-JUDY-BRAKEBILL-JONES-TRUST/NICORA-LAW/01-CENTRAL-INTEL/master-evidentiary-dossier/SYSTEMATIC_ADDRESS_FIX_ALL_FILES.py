#!/usr/bin/env python3
"""
SYSTEMATIC ADDRESS FIX: Fix Eric's address in ALL documents
Following PFV v5.1 systematic approach
"""

from docx import Document
import os
import glob

def fix_address_in_document(doc):
    """Fix Eric's personal address in all locations"""

    fixes = []

    # Process all paragraphs
    for idx, para in enumerate(doc.paragraphs):
        original_text = para.text

        # Fix personal address (appears in header, not property descriptions)
        if "17742 Berta Canyon Road" in para.text:
            # Check if this is Eric's personal address context (header area)
            if idx < 15 or "Tel:" in para.text or "Email:" in para.text or "IN PRO PER" in para.text:
                para.text = para.text.replace("17742 Berta Canyon Road", "5634 Noel Drive")
                fixes.append(f"Para {idx}: Changed address to 5634 Noel Drive")

        if "Lake Hughes, CA 93532" in para.text:
            para.text = para.text.replace("Lake Hughes, CA 93532", "Temple City, CA 91780")
            fixes.append(f"Para {idx}: Changed city to Temple City")

        if "Lake Hughes" in para.text and idx < 15:
            para.text = para.text.replace("Lake Hughes", "Temple City")
            fixes.append(f"Para {idx}: Changed Lake Hughes → Temple City")

        if "93532" in para.text and idx < 15:
            para.text = para.text.replace("93532", "91780")
            fixes.append(f"Para {idx}: Changed ZIP 93532 → 91780")

    # Process all tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    original_text = para.text

                    # Fix address in tables (usually header caption)
                    if "17742 Berta Canyon Road" in para.text:
                        para.text = para.text.replace("17742 Berta Canyon Road", "5634 Noel Drive")
                        fixes.append(f"Table {table_idx} Row {row_idx}: Changed address")

                    if "Lake Hughes" in para.text:
                        para.text = para.text.replace("Lake Hughes, CA 93532", "Temple City, CA 91780")
                        para.text = para.text.replace("Lake Hughes", "Temple City")
                        fixes.append(f"Table {table_idx} Row {row_idx}: Changed city")

                    if "93532" in para.text and "17742" not in para.text:  # Don't change property ZIP
                        para.text = para.text.replace("93532", "91780")
                        fixes.append(f"Table {table_idx} Row {row_idx}: Changed ZIP")

    return doc, fixes

def convert_to_pdf(docx_path):
    """Convert DOCX to PDF"""
    output_dir = os.path.dirname(docx_path)
    basename = os.path.basename(docx_path)

    cmd = f'cd "{output_dir}" && /Applications/LibreOffice.app/Contents/MacOS/soffice --headless --convert-to pdf --outdir . "{basename}" 2>&1 > /dev/null'
    result = os.system(cmd)

    pdf_path = docx_path.replace('.docx', '.pdf')
    return os.path.exists(pdf_path)

def main():
    print("\n" + "="*70)
    print("SYSTEMATIC ADDRESS FIX - ALL FILES")
    print("="*70)
    print("\n✓ CORRECT ADDRESS:")
    print("  Eric Brakebill Jones")
    print("  5634 Noel Drive")
    print("  Temple City, CA 91780")
    print("  (626) 348-3019")
    print("  eric@recovery-compass.org")
    print("\n✓ PROPERTY ADDRESS (Unchanged):")
    print("  17742 Berta Canyon Road")
    print("  Salinas, CA 93907")
    print("\n" + "="*70 + "\n")

    output_dir = "/Users/ericjones/Compass Fortress/cases/judy-jones-probate-850/4-email-to-anuar-nov15/ALL-ATTACHMENTS-READY-TO-SEND"

    # Find ALL DOCX files
    docx_files = glob.glob(os.path.join(output_dir, "*.docx"))
    docx_files.sort()

    print(f"Found {len(docx_files)} DOCX files to process\n")

    results = {
        'fixed': [],
        'no_changes': [],
        'failed': []
    }

    for docx_path in docx_files:
        filename = os.path.basename(docx_path)

        print(f"{'='*70}")
        print(f"Processing: {filename}")
        print(f"{'='*70}")

        try:
            # Load document
            doc = Document(docx_path)
            print(f"✓ Loaded")

            # Fix address
            doc, fixes = fix_address_in_document(doc)

            if fixes:
                print(f"✓ Made {len(fixes)} corrections:")
                for fix in fixes[:5]:  # Show first 5
                    print(f"  • {fix}")
                if len(fixes) > 5:
                    print(f"  ... and {len(fixes) - 5} more")

                # Save
                doc.save(docx_path)
                print(f"✓ Saved")

                # Convert to PDF
                if convert_to_pdf(docx_path):
                    print(f"✓ PDF regenerated")
                    results['fixed'].append(filename)
                else:
                    print(f"⚠ PDF conversion may have failed")
                    results['fixed'].append(filename)

            else:
                print(f"• No changes needed (may already be correct)")
                results['no_changes'].append(filename)

        except Exception as e:
            print(f"✗ ERROR: {e}")
            results['failed'].append(filename)

        print()

    # Summary
    print("="*70)
    print("FINAL SUMMARY")
    print("="*70)
    print(f"\n✓ Fixed: {len(results['fixed'])} files")
    if results['fixed']:
        for f in results['fixed']:
            print(f"  • {f}")

    print(f"\n• No changes: {len(results['no_changes'])} files")
    if results['no_changes']:
        for f in results['no_changes'][:5]:
            print(f"  • {f}")

    if results['failed']:
        print(f"\n✗ Failed: {len(results['failed'])} files")
        for f in results['failed']:
            print(f"  • {f}")

    print("\n" + "="*70)
    print("✅ ADDRESS FIX COMPLETE")
    print("="*70)
    print("\nAll documents should now show:")
    print("  5634 Noel Drive, Temple City, CA 91780")
    print("\nVerify by opening any PDF and checking header.")
    print("="*70 + "\n")

if __name__ == "__main__":
    main()
