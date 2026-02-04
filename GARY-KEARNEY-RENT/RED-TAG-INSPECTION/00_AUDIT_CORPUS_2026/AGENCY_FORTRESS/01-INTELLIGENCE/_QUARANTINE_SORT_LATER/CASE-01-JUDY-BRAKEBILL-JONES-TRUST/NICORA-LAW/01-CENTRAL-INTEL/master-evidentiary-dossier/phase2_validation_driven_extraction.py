#!/usr/bin/env python3
"""
Phase 2: Validation-Driven Content Extraction with Audit Trail
This script extracts content intelligently and validates what was extracted
"""

from docx import Document
import os
import re

# ===== EXPECTED CONTENT VALIDATION RULES =====

EXPECTED_CONTENT = {
    '01_Ex_Parte_Application': {
        'required_sections': [
            'STATEMENT OF FACTS',
            'LEGAL ARGUMENT',
            'CONCLUSION'
        ],
        'required_phrases': [
            'Sole Successor Trustee',
            'foreclosure',
            'December 3',
            'emergency',
            'ex parte',
            'good cause'
        ],
        'min_paragraphs': 8
    },
    '02_Declaration_with_Exhibits': {
        'required_sections': [],
        'required_phrases': [
            'I am over the age of 18',
            'competent to testify',
            'The Trustor',
            'died on April 4, 2025',
            'Respondents',
            'Mercedes',
            '$10,650',
            'Chase Bank',
            'deposit slip',
            'I declare under penalty of perjury'
        ],
        'min_paragraphs': 12,
        'numbered_paragraphs': True
    },
    '04_MPA': {
        'required_sections': [
            'INTRODUCTION',
            'FACTUAL BACKGROUND',
            'LEGAL STANDARD',
            'ARGUMENT'
        ],
        'required_phrases': [
            '§ 859',
            'double damages',
            'bad faith',
            'RUFADAA',
            'digital asset',
            'irreparable harm'
        ],
        'min_paragraphs': 15
    },
    '06_Petition_850': {
        'required_sections': [
            'JURISDICTION',
            'PARTIES'
        ],
        'required_phrases': [
            'Probate Code',
            '§ 850',
            'Heggstad',
            'Schedule A',
            'successor trustee'
        ],
        'min_paragraphs': 8
    }
}

# ===== CONTENT EXTRACTION FUNCTIONS =====

def is_numbered_paragraph(text):
    """Check if paragraph starts with a number (1., 2., 3., etc.)"""
    text = text.strip()
    if len(text) < 3:
        return False

    # Match patterns like "1.", "2.", "14.", etc. at start
    match = re.match(r'^(\d+)\.\s+', text)
    return match is not None

def is_section_heading(text):
    """Check if paragraph is a section heading (I., II., A., B., etc.)"""
    text = text.strip().upper()

    patterns = [
        r'^I\.\s+',           # I.
        r'^II\.\s+',          # II.
        r'^III\.\s+',         # III.
        r'^IV\.\s+',          # IV.
        r'^A\.\s+',           # A.
        r'^B\.\s+',           # B.
        r'^C\.\s+',           # C.
    ]

    for pattern in patterns:
        if re.match(pattern, text):
            return True

    # Also check for common section titles
    section_titles = [
        'INTRODUCTION',
        'STATEMENT OF FACTS',
        'FACTUAL BACKGROUND',
        'LEGAL STANDARD',
        'LEGAL ARGUMENT',
        'ARGUMENT',
        'CONCLUSION',
        'JURISDICTION',
        'PARTIES'
    ]

    for title in section_titles:
        if title in text:
            return True

    return False

def is_likely_exhibit(text):
    """Check if content is likely an embedded exhibit"""
    text = text.lower()

    exhibit_indicators = [
        'exhibit',
        'attachment',
        'see attached',
        'nicora law',
        'attorney at law',
        'law office of',
        'facsimile',
        'fax:',
        're:',
        'dear '
    ]

    # If text contains multiple exhibit indicators, likely an exhibit
    indicator_count = sum(1 for indicator in exhibit_indicators if indicator in text)

    return indicator_count >= 2

def is_substantive_content(text):
    """Check if paragraph contains substantive legal content"""
    text = text.strip()

    # Too short
    if len(text) < 30:
        return False

    # Header-like
    if any(keyword in text for keyword in ['SUPERIOR COURT', 'COUNTY OF', 'CASE NO.', 'Tel:', 'Email:']):
        return False

    # Likely exhibit
    if is_likely_exhibit(text):
        return False

    return True

def extract_content_intelligently(source_doc, doc_name):
    """Extract content with intelligent filtering"""

    extracted = {
        'numbered_paragraphs': [],
        'section_headings': [],
        'substantive_content': [],
        'all_extracted': []
    }

    print(f"\n  📖 Reading source document...")
    print(f"     Total paragraphs in source: {len(source_doc.paragraphs)}")

    for i, para in enumerate(source_doc.paragraphs):
        text = para.text.strip()

        # Skip empty
        if len(text) < 10:
            continue

        # Skip header region (first 15 paragraphs)
        if i < 15:
            continue

        # Check what type of content this is
        if is_numbered_paragraph(text):
            extracted['numbered_paragraphs'].append((i, para, text[:80]))
            extracted['all_extracted'].append(para)
            print(f"     ✓ Line {i}: Numbered paragraph found: {text[:60]}...")

        elif is_section_heading(text):
            extracted['section_headings'].append((i, para, text[:80]))
            extracted['all_extracted'].append(para)
            print(f"     ✓ Line {i}: Section heading found: {text[:60]}...")

        elif is_substantive_content(text) and not is_likely_exhibit(text):
            extracted['substantive_content'].append((i, para, text[:80]))
            extracted['all_extracted'].append(para)
            print(f"     ✓ Line {i}: Substantive content: {text[:60]}...")

        else:
            if 'nicora' in text.lower() or 'exhibit' in text.lower():
                print(f"     ⊗ Line {i}: Skipped (likely exhibit): {text[:60]}...")
            elif len(text) >= 20:
                print(f"     ⊗ Line {i}: Skipped (filtered): {text[:40]}...")

    print(f"\n  📊 Extraction Summary:")
    print(f"     Numbered paragraphs: {len(extracted['numbered_paragraphs'])}")
    print(f"     Section headings: {len(extracted['section_headings'])}")
    print(f"     Substantive content: {len(extracted['substantive_content'])}")
    print(f"     Total extracted: {len(extracted['all_extracted'])}")

    return extracted

def validate_extracted_content(extracted, doc_name):
    """Validate that extracted content matches expectations"""

    print(f"\n  🔍 Validating extracted content...")

    if doc_name not in EXPECTED_CONTENT:
        print(f"     ⚠ No validation rules for {doc_name}")
        return {'passed': False, 'failures': ['No validation rules defined']}

    rules = EXPECTED_CONTENT[doc_name]
    failures = []
    warnings = []

    # Build full text from extracted content
    full_text = ' '.join([para.text for para in extracted['all_extracted']])

    # Check required sections
    for section in rules.get('required_sections', []):
        if section not in full_text:
            failures.append(f"Missing required section: {section}")
            print(f"     ✗ Missing section: {section}")
        else:
            print(f"     ✓ Found section: {section}")

    # Check required phrases
    found_phrases = 0
    total_phrases = len(rules.get('required_phrases', []))

    for phrase in rules.get('required_phrases', []):
        if phrase.lower() in full_text.lower():
            found_phrases += 1
            print(f"     ✓ Found phrase: '{phrase}'")
        else:
            warnings.append(f"Missing phrase: {phrase}")
            print(f"     ⚠ Missing phrase: '{phrase}'")

    # Check minimum paragraphs
    min_paras = rules.get('min_paragraphs', 0)
    actual_paras = len(extracted['all_extracted'])

    if actual_paras < min_paras:
        failures.append(f"Insufficient paragraphs: {actual_paras} < {min_paras} expected")
        print(f"     ✗ Only {actual_paras} paragraphs (need {min_paras})")
    else:
        print(f"     ✓ Paragraph count OK: {actual_paras} >= {min_paras}")

    # Check numbered paragraphs for declarations
    if rules.get('numbered_paragraphs'):
        num_count = len(extracted['numbered_paragraphs'])
        if num_count < 10:
            failures.append(f"Insufficient numbered paragraphs: {num_count} < 10 expected")
            print(f"     ✗ Only {num_count} numbered paragraphs (need ~12-14)")
        else:
            print(f"     ✓ Numbered paragraphs OK: {num_count}")

    # Calculate phrase coverage
    phrase_coverage = (found_phrases / total_phrases * 100) if total_phrases > 0 else 0
    print(f"\n  📊 Validation Score:")
    print(f"     Phrase coverage: {phrase_coverage:.0f}% ({found_phrases}/{total_phrases})")
    print(f"     Failures: {len(failures)}")
    print(f"     Warnings: {len(warnings)}")

    passed = len(failures) == 0 and phrase_coverage >= 60

    return {
        'passed': passed,
        'failures': failures,
        'warnings': warnings,
        'phrase_coverage': phrase_coverage
    }

def merge_with_validation(template_path, source_path, output_path, eric_info, doc_name):
    """Merge with comprehensive validation"""

    print(f"\n{'='*70}")
    print(f"Processing: {doc_name}")
    print(f"{'='*70}")

    # Load template
    template = Document(template_path)
    print(f"\n✓ Loaded template: {len(template.paragraphs)} paras, {len(template.tables)} tables")

    # Update placeholders in template
    print(f"\n✓ Updating template placeholders...")
    template = update_placeholders_everywhere(template, eric_info)

    # Load source
    try:
        source = Document(source_path)
        print(f"\n✓ Loaded source document")
    except Exception as e:
        print(f"\n✗ Error loading source: {e}")
        return None

    # Extract content intelligently
    extracted = extract_content_intelligently(source, doc_name)

    # Validate extracted content
    validation = validate_extracted_content(extracted, doc_name)

    # Build merged document
    print(f"\n  🔨 Building merged document...")

    # Keep header (first 16 paragraphs)
    while len(template.paragraphs) > 16:
        p = template.paragraphs[-1]
        p._element.getparent().remove(p._element)

    # Remove signature table
    if len(template.tables) > 1:
        tbl = template.tables[-1]._element
        tbl.getparent().remove(tbl)

    template.add_paragraph("")

    # Insert extracted content
    added = 0
    for para in extracted['all_extracted']:
        if len(para.text.strip()) > 0:
            new_para = template.add_paragraph(para.text)

            # Preserve formatting
            if para.runs and new_para.runs:
                for src_run, dst_run in zip(para.runs[:len(new_para.runs)], new_para.runs):
                    if src_run.bold:
                        dst_run.bold = True
                    if src_run.italic:
                        dst_run.italic = True
                    if src_run.underline:
                        dst_run.underline = True

            added += 1

    print(f"     Added {added} paragraphs to template")

    # Add signature block
    template.add_paragraph("")
    template.add_paragraph("")
    template.add_paragraph("DATED: November 15, 2025")
    template.add_paragraph("")
    template.add_paragraph("Respectfully submitted,")
    template.add_paragraph("")
    template.add_paragraph("_______________________________")
    template.add_paragraph("ERIC BRAKEBILL JONES")
    template.add_paragraph("Sole Successor Trustee, In Pro Per")
    template.add_paragraph("Judy Brakebill Jones 2008 Revocable Trust")

    # Save
    template.save(output_path)
    print(f"\n✓ Saved: {os.path.basename(output_path)}")

    return validation

def update_placeholders_everywhere(doc, eric_info):
    """Update all placeholders"""
    replacements_made = 0

    for para in doc.paragraphs:
        original = para.text
        para.text = apply_replacements(para.text, eric_info)
        if original != para.text:
            replacements_made += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    original = para.text
                    para.text = apply_replacements(para.text, eric_info)
                    if original != para.text:
                        replacements_made += 1

    print(f"  • Made {replacements_made} placeholder replacements")
    return doc

def apply_replacements(text, eric_info):
    """Apply all replacements"""
    text = text.replace("ADDRESS", eric_info['address'])
    text = text.replace("CITY, CA ZIP", eric_info['city_state_zip'])
    text = text.replace("(415) 123-4567", eric_info['phone'])
    text = text.replace("Tel: (415) 123-4567", f"Tel: {eric_info['phone']}")
    text = text.replace("email address", eric_info['email'])
    text = text.replace("Email: email address", f"Email: {eric_info['email']}")
    text = text.replace("LARRY LASAGNA 2002 REVOCABLE TRUST", eric_info['trust'].upper())
    text = text.replace("LARRY LASAGNA", "JUDY BRAKEBILL JONES")
    text = text.replace("JUDY [SURNMAE]", "JUDY BRAKEBILL JONES")
    text = text.replace("[SURNMAE]", "BRAKEBILL JONES")
    text = text.replace("ESTATE OF LARRY LASAGNA", "ESTATE OF JUDY BRAKEBILL JONES")
    text = text.replace("[NAME COUNTY]", "LOS ANGELES")
    text = text.replace("COUNTY OF [NAME COUNTY]", "COUNTY OF LOS ANGELES")
    text = text.replace("CASE NO. RP21094431", "CASE NO.: (To be assigned)")
    text = text.replace("RP21094431", "(To be assigned)")
    text = text.replace('NAME OF THE DOCUMENT FOR EXAMPLE "EX PARTE APPLICATION REQUE', 'EX PARTE APPLICATION')
    text = text.replace("LAW OFFICE OF DAVID W. TATE", "")
    text = text.replace("DAVID W. TATE, ESQ.", "")
    text = text.replace("Attorneys for all Respondents and", "")
    text = text.replace("Moving Parties", "")
    text = text.replace("____________, 2023", "November 15, 2025")
    text = text.replace("February 9, 2024", "November 15, 2025")
    text = text.replace("July 13, 2021", "November 15, 2025")
    text = text.replace("Reservation ID: 780799297986", "")
    text = text.replace("Time: 10:30 AM", "")
    text = text.replace("Dept.: Probate, 201", "")
    text = text.replace("Action Filed: July 13, 2021", "")
    return text

def convert_to_pdf(docx_path):
    """Convert to PDF"""
    output_dir = os.path.dirname(docx_path)
    basename = os.path.basename(docx_path)

    cmd = f'cd "{output_dir}" && /Applications/LibreOffice.app/Contents/MacOS/soffice --headless --convert-to pdf --outdir . "{basename}" 2>&1 > /dev/null'
    result = os.system(cmd)

    pdf_path = docx_path.replace('.docx', '.pdf')
    if os.path.exists(pdf_path):
        size_kb = os.path.getsize(pdf_path) / 1024
        print(f"✓ PDF: {os.path.basename(pdf_path)} ({size_kb:.1f} KB)")
        return True
    return False

# ===== CONFIGURATION =====

ERIC_INFO = {
    'name': 'ERIC BRAKEBILL JONES',
    'address': '17742 Berta Canyon Road',
    'city_state_zip': 'Lake Hughes, CA 93532',
    'phone': '(626) 348-3019',
    'email': 'eric@recovery-compass.org',
    'capacity': 'Sole Successor Trustee, In Pro Per',
    'trust': 'Judy Brakebill Jones 2008 Revocable Trust'
}

TEMPLATE_PATH = "/Users/ericjones/Cases/Judy-Trust/pro-per-probate-850/anuar-medina-ramirez-seven-hills-law/email-november-11/pleading-paper/PLEADING PAPER FOR ERIC JONES.docx"

SOURCE_DIR = "/Users/ericjones/11-15/5-bird/judy-jones-probate-850/self-filing-IRONCLAD-FINAL/remediated_docx_v2"

OUTPUT_DIR = "/Users/ericjones/Compass Fortress/cases/judy-jones-probate-850/4-email-to-anuar-nov15/ALL-ATTACHMENTS-READY-TO-SEND"

DOCUMENTS = [
    ('01_Ex_Parte_Application', '01_Ex_Parte_Application.docx', '01_Ex_Parte_Application_CRC2111_VALIDATED.docx'),
    ('02_Declaration_with_Exhibits', '02_Declaration_with_Exhibits.docx', '02_Declaration_with_Exhibits_CRC2111_VALIDATED.docx'),
    ('04_MPA', '04_MPA.docx', '04_MPA_CRC2111_VALIDATED.docx'),
    ('06_Petition_850', '06_Petition_850.docx', '06_Petition_850_CRC2111_VALIDATED.docx')
]

# ===== MAIN =====

def main():
    print("\n" + "="*70)
    print("PHASE 2: VALIDATION-DRIVEN CONTENT EXTRACTION")
    print("="*70)
    print("\nThis script validates extracted content against expected requirements.")
    print("Honest reporting: Will identify gaps for Anuar to complete.\n")

    results = []

    for doc_name, source_file, output_file in DOCUMENTS:
        source_path = os.path.join(SOURCE_DIR, source_file)
        output_path = os.path.join(OUTPUT_DIR, output_file)

        if not os.path.exists(source_path):
            print(f"\n✗ Source missing: {source_file}")
            results.append({
                'name': doc_name,
                'status': 'FAILED',
                'reason': 'Source file not found'
            })
            continue

        # Merge with validation
        validation = merge_with_validation(TEMPLATE_PATH, source_path, output_path, ERIC_INFO, doc_name)

        if validation is None:
            results.append({
                'name': doc_name,
                'status': 'FAILED',
                'reason': 'Error during processing'
            })
            continue

        # Convert to PDF
        pdf_success = convert_to_pdf(output_path)

        # Record result
        if validation['passed']:
            status = 'COMPLETE'
        elif len(validation['failures']) == 0 and validation['phrase_coverage'] >= 50:
            status = 'PARTIAL'
        else:
            status = 'INCOMPLETE'

        results.append({
            'name': doc_name,
            'status': status,
            'validation': validation,
            'pdf_created': pdf_success
        })

    # ===== AUDIT REPORT =====
    print("\n" + "="*70)
    print("AUDIT REPORT: HONEST ASSESSMENT")
    print("="*70)

    complete_count = sum(1 for r in results if r['status'] == 'COMPLETE')
    partial_count = sum(1 for r in results if r['status'] == 'PARTIAL')
    incomplete_count = sum(1 for r in results if r['status'] == 'INCOMPLETE')

    print(f"\n📊 Overall Status:")
    print(f"   ✓ Complete: {complete_count}/4")
    print(f"   ⚠ Partial: {partial_count}/4")
    print(f"   ✗ Incomplete: {incomplete_count}/4")

    print(f"\n📄 Document-by-Document Results:\n")

    for result in results:
        print(f"{'='*70}")
        print(f"{result['name']}")
        print(f"Status: {result['status']}")

        if 'validation' in result:
            val = result['validation']
            print(f"Phrase Coverage: {val['phrase_coverage']:.0f}%")

            if val['failures']:
                print(f"\n✗ Validation Failures:")
                for failure in val['failures']:
                    print(f"   • {failure}")

            if val['warnings']:
                print(f"\n⚠ Warnings:")
                for warning in val['warnings'][:5]:  # Show first 5
                    print(f"   • {warning}")

        print()

    # ===== RECOMMENDATIONS =====
    print("="*70)
    print("RECOMMENDATIONS FOR ANUAR")
    print("="*70)

    for result in results:
        if result['status'] != 'COMPLETE':
            print(f"\n{result['name']}:")

            if 'validation' in result:
                val = result['validation']

                if val['failures']:
                    print("  Missing Critical Content:")
                    for failure in val['failures']:
                        print(f"    • {failure}")

                if val['phrase_coverage'] < 80:
                    print(f"  Content Coverage: {val['phrase_coverage']:.0f}% (needs manual review)")

    print("\n" + "="*70)
    print("✉️  FILES READY FOR ANUAR REVIEW")
    print("="*70)
    print(f"\nLocation: {OUTPUT_DIR}/")
    print("\n_VALIDATED.pdf files contain best-effort extraction.")
    print("Review needed for completeness against source documents.")
    print("\n" + "="*70 + "\n")

if __name__ == "__main__":
    main()
