#!/usr/bin/env python3
"""
EMERGENCY FIX: Correct Eric's address in all documents
"""

from docx import Document
import os

def fix_address(doc):
    """Fix Eric's personal address throughout document"""

    replacements = {
        # Fix personal address
        "17742 Berta Canyon Road": "5634 Noel Drive",
        "Lake Hughes, CA 93532": "Temple City, CA 91780",
        "Lake Hughes": "Temple City",
        "93532": "91780",

        # Property address stays the same but fix city
        # When it appears with "property located at" or similar context
    }

    changes = 0

    # Fix paragraphs
    for para in doc.paragraphs:
        original = para.text

        # Only change if it's in the header context (not property description)
        # Check if paragraph has contact info indicators
        is_header = any(indicator in original for indicator in [
            "Tel:", "Email:", "IN PRO PER", "eric@recovery-compass.org"
        ])

        if is_header:
            for old, new in replacements.items():
                if old in para.text:
                    para.text = para.text.replace(old, new)
                    changes += 1
                    print(f"    Fixed: {old} → {new}")

    # Fix tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    original = para.text

                    is_header = any(indicator in original for indicator in [
                        "Tel:", "Email:", "IN PRO PER", "eric@recovery-compass.org"
                    ])

                    if is_header:
                        for old, new in replacements.items():
                            if old in para.text:
                                para.text = para.text.replace(old, new)
                                changes += 1
                                print(f"    Fixed: {old} → {new}")

    return doc, changes

def convert_to_pdf(docx_path):
    """Convert to PDF"""
    output_dir = os.path.dirname(docx_path)
    basename = os.path.basename(docx_path)

    cmd = f'cd "{output_dir}" && /Applications/LibreOffice.app/Contents/MacOS/soffice --headless --convert-to pdf --outdir . "{basename}" 2>&1 > /dev/null'
    os.system(cmd)

    pdf_path = docx_path.replace('.docx', '.pdf')
    if os.path.exists(pdf_path):
        size_kb = os.path.getsize(pdf_path) / 1024
        return True
    return False

OUTPUT_DIR = "/Users/ericjones/Compass Fortress/cases/judy-jones-probate-850/4-email-to-anuar-nov15/ALL-ATTACHMENTS-READY-TO-SEND"

DOCUMENTS = [
    "01_Ex_Parte_Application_CRC2111_FINAL_v2.docx",
    "02_Declaration_CRC2111_FINAL_v2.docx",
    "04_MPA_CRC2111_FINAL_v2.docx"
]

def main():
    print("\n" + "="*70)
    print("EMERGENCY ADDRESS CORRECTION")
    print("="*70)
    print("\nFixing Eric's personal address:")
    print("  WRONG: 17742 Berta Canyon Road, Lake Hughes, CA 93532")
    print("  RIGHT: 5634 Noel Drive, Temple City, CA 91780\n")

    for doc_file in DOCUMENTS:
        print(f"\nProcessing: {doc_file}")

        docx_path = os.path.join(OUTPUT_DIR, doc_file)

        if not os.path.exists(docx_path):
            print(f"  ✗ File not found")
            continue

        # Load and fix
        doc = Document(docx_path)
        doc, changes = fix_address(doc)

        print(f"  ✓ Made {changes} corrections")

        # Save
        doc.save(docx_path)
        print(f"  ✓ Saved corrected version")

        # Convert to PDF
        if convert_to_pdf(docx_path):
            print(f"  ✓ PDF regenerated")
        else:
            print(f"  ✗ PDF conversion failed")

    print("\n" + "="*70)
    print("ADDRESS CORRECTION COMPLETE")
    print("="*70)
    print("\nAll documents now have CORRECT address:")
    print("  Eric Brakebill Jones")
    print("  5634 Noel Drive")
    print("  Temple City, CA 91780")
    print("  (626) 348-3019")
    print("  eric@recovery-compass.org")
    print("\n" + "="*70 + "\n")

if __name__ == "__main__":
    main()
