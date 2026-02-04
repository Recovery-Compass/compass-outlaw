#!/usr/bin/env python3
"""
FIX CRITICAL FAILURE: Replace ALL template placeholder data with Eric's real information
"""

from docx import Document
from pathlib import Path
import subprocess

print("=" * 70)
print("CRITICAL FIX: Replacing Template Dummy Data with Real Case Info")
print("=" * 70)

BASE_DIR = Path("/Users/ericjones/Compass Fortress/cases/judy-jones-probate-850")
SOURCE_DIR = BASE_DIR / "1-core-filing-documents/PFV_V12_ATTORNEY_READY"
OUTPUT_DIR = BASE_DIR / "1-core-filing-documents/CORRECTED_FINAL"

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Eric's REAL information (from the formatted text files he provided)
REAL_INFO = {
    # Eric's contact info
    "ERIC B. JONES (IN PRO PER)": "ERIC BRAKEBILL JONES",
    "ADDRESS": "17742 Berta Canyon Road",
    "CITY, CA ZIP": "Lake Hughes, CA 93532",
    "Tel: (415) 123-4567": "Phone: (626) 348-3019",
    "Email: email address": "Email: eric@recovery-compass.org",

    # Trustee/Capacity
    "In pro per, as Trustee of the LARRY LASAGNA 2002 REVOCABLE TRUST,": "Sole Successor Trustee",
    "and as Executor of the ESTATE OF LARRY LASAGNA, Deceased,": "In Pro Per",
    "LARRY LASAGNA 2002 REVOCABLE TRUST": "Judy Brakebill Jones 2008 Revocable Trust",
    "LARRY LASAGNA": "Judy Brakebill Jones",

    # Court info
    "COUNTY OF [NAME COUNTY]": "COUNTY OF LOS ANGELES - PROBATE DIVISION",
    "[NAME COUNTY]": "LOS ANGELES",
    "COUNTY OF MONTEREY": "COUNTY OF LOS ANGELES",
    "MONTEREY": "LOS ANGELES",

    # Case info
    "IN RE THE ESTATE OF JUDY [SURNMAE]": "In the Matter of: The Judy Brakebill Jones 2008 Revocable Trust",
    "[SURNMAE]": "BRAKEBILL JONES",
    "CASE NO. RP21094431": "Case No.: (to be assigned)",
    "RP21094431": "(to be assigned)",
    "[TO BE ASSIGNED]": "(to be assigned)",

    # Dates
    "Date: February 9, 2024": f"Date: November 15, 2025",
    "February 9, 2024": "November 15, 2025",
    ", 2023": ", 2025",
    "DATED: , 2023": f"Executed on November 15, 2025, at Lake Hughes, California.",

    # Attorney info (remove - Eric is pro per)
    "LAW OFFICE OF DAVID W. TATE": "",
    "DAVID W. TATE": "",

    # Department/court details
    "Dept.: Probate, 201": "",
    "Action Filed: July 13, 2021": "",
    "Reservation ID: 780799297986": "",
    "Time: 10:30 AM": "",
}

def fix_document_placeholders(input_path, output_path):
    """Replace ALL template placeholder data with real case information."""

    print(f"\n📋 Fixing: {input_path.name}")

    doc = Document(str(input_path))
    replacements_made = 0

    # Fix all paragraphs
    for para in doc.paragraphs:
        original_text = para.text
        new_text = original_text

        for placeholder, real_value in REAL_INFO.items():
            if placeholder in new_text:
                new_text = new_text.replace(placeholder, real_value)
                replacements_made += 1

        if new_text != original_text:
            para.text = new_text

    # Fix all tables (including header table)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    original_text = para.text
                    new_text = original_text

                    for placeholder, real_value in REAL_INFO.items():
                        if placeholder in new_text:
                            new_text = new_text.replace(placeholder, real_value)
                            replacements_made += 1

                    if new_text != original_text:
                        para.text = new_text

    # Fix headers
    for section in doc.sections:
        header = section.header
        for para in header.paragraphs:
            original_text = para.text
            new_text = original_text

            for placeholder, real_value in REAL_INFO.items():
                if placeholder in new_text:
                    new_text = new_text.replace(placeholder, real_value)

            if new_text != original_text:
                para.text = new_text

    # Fix footers
    for section in doc.sections:
        footer = section.footer
        for para in footer.paragraphs:
            original_text = para.text
            new_text = original_text

            for placeholder, real_value in REAL_INFO.items():
                if placeholder in new_text:
                    new_text = new_text.replace(placeholder, real_value)

            if new_text != original_text:
                para.text = new_text

    doc.save(str(output_path))
    print(f"   ✅ Fixed {replacements_made} placeholder instances")
    return output_path

def convert_to_pdf(docx_path, pdf_path):
    """Convert to PDF using LibreOffice."""
    try:
        subprocess.run([
            "soffice", "--headless", "--convert-to", "pdf",
            "--outdir", str(pdf_path.parent), str(docx_path)
        ], check=True, capture_output=True)

        expected_pdf = docx_path.parent / f"{docx_path.stem}.pdf"
        if expected_pdf.exists() and expected_pdf != pdf_path:
            expected_pdf.rename(pdf_path)

        if pdf_path.exists():
            size_kb = pdf_path.stat().st_size / 1024
            print(f"   ✅ PDF: {pdf_path.name} ({size_kb:.1f} KB)")
            return True
    except Exception as e:
        print(f"   ❌ PDF failed: {e}")
        return False

# Process all documents
documents = [
    "01_Ex_Parte_Application_CRC2111_FINAL.docx",
    "02_Declaration_with_Exhibits_CRC2111_FINAL.docx",
    "04_MPA_CRC2111_FINAL.docx",
    "06_Petition_850_CRC2111_FINAL.docx"
]

success = 0
for doc_name in documents:
    input_path = SOURCE_DIR / doc_name
    output_docx = OUTPUT_DIR / doc_name.replace("_FINAL", "_CORRECTED")
    output_pdf = output_docx.with_suffix('.pdf')

    if input_path.exists():
        fix_document_placeholders(input_path, output_docx)
        if convert_to_pdf(output_docx, output_pdf):
            success += 1
    else:
        print(f"❌ Not found: {doc_name}")

print("\n" + "=" * 70)
print(f"✅ CORRECTED: {success}/4 documents with real case information")
print(f"📁 Output: {OUTPUT_DIR}")
print("=" * 70)
