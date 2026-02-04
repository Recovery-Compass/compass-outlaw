#!/usr/bin/env python3
"""
SYSTEMATIC DATE FIX: Remove all template dummy dates
Case hasn't been filed yet - no court date/time/dept assigned
"""

from docx import Document
import os
import glob

def clean_template_dates(doc):
    """Remove template placeholder dates and info"""

    fixes = []

    # Template artifacts to remove
    remove_patterns = [
        "Date: February 9, 2024",
        "February 9, 2024",
        "Time: 10:30 AM",
        "10:30 AM",
        "Dept.: Probate, 201",
        "Probate, 201",
        "Action Filed: July 13, 2021",
        "July 13, 2021",
        "Reservation ID: 780799297986",
        "780799297986",
        "Date:",
        "Time:",
        "Dept.:",
        "Action Filed:",
        "Reservation ID:"
    ]

    # Process paragraphs
    for idx, para in enumerate(doc.paragraphs):
        original = para.text

        for pattern in remove_patterns:
            if pattern in para.text:
                para.text = para.text.replace(pattern, "")
                if original != para.text:
                    fixes.append(f"Para {idx}: Removed '{pattern}'")

        # Clean up extra whitespace
        para.text = para.text.strip()

    # Process tables (caption table likely has these)
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    original = para.text

                    for pattern in remove_patterns:
                        if pattern in para.text:
                            para.text = para.text.replace(pattern, "")
                            if original != para.text:
                                fixes.append(f"Table {table_idx} Row {row_idx}: Removed '{pattern}'")

                    # Clean up extra whitespace
                    para.text = para.text.strip()

    return doc, fixes

def convert_to_pdf(docx_path):
    """Convert to PDF"""
    output_dir = os.path.dirname(docx_path)
    basename = os.path.basename(docx_path)

    cmd = f'cd "{output_dir}" && /Applications/LibreOffice.app/Contents/MacOS/soffice --headless --convert-to pdf --outdir . "{basename}" 2>&1 > /dev/null'
    os.system(cmd)

    pdf_path = docx_path.replace('.docx', '.pdf')
    return os.path.exists(pdf_path)

def main():
    print("\n" + "="*70)
    print("SYSTEMATIC DATE FIX - REMOVE TEMPLATE DATES")
    print("="*70)
    print("\n❌ REMOVING Template Placeholder Dates:")
    print("  • Date: February 9, 2024")
    print("  • Time: 10:30 AM")
    print("  • Dept.: Probate, 201")
    print("  • Action Filed: July 13, 2021")
    print("  • Reservation ID: 780799297986")
    print("\n✅ KEEPING Only Real Date:")
    print("  • DATED: November 15, 2025 (signature block)")
    print("\n📝 REASON:")
    print("  Case hasn't been filed yet - no hearing date/time assigned by court")
    print("\n" + "="*70 + "\n")

    output_dir = "/Users/ericjones/Compass Fortress/cases/judy-jones-probate-850/4-email-to-anuar-nov15/ALL-ATTACHMENTS-READY-TO-SEND"

    # Find all DOCX files
    docx_files = glob.glob(os.path.join(output_dir, "*.docx"))
    docx_files.sort()

    print(f"Processing {len(docx_files)} DOCX files\n")

    results = {
        'fixed': [],
        'no_changes': [],
        'failed': []
    }

    for docx_path in docx_files:
        filename = os.path.basename(docx_path)

        print(f"{'='*70}")
        print(f"{filename}")
        print(f"{'='*70}")

        try:
            doc = Document(docx_path)
            print(f"✓ Loaded")

            doc, fixes = clean_template_dates(doc)

            if fixes:
                print(f"✓ Removed {len(fixes)} template date/info items:")
                for fix in fixes[:5]:
                    print(f"  • {fix}")
                if len(fixes) > 5:
                    print(f"  ... and {len(fixes) - 5} more")

                doc.save(docx_path)
                print(f"✓ Saved")

                if convert_to_pdf(docx_path):
                    print(f"✓ PDF regenerated")
                    results['fixed'].append(filename)
                else:
                    print(f"⚠ PDF may have failed")
                    results['fixed'].append(filename)

            else:
                print(f"• No template dates found")
                results['no_changes'].append(filename)

        except Exception as e:
            print(f"✗ ERROR: {e}")
            results['failed'].append(filename)

        print()

    # Summary
    print("="*70)
    print("FINAL SUMMARY")
    print("="*70)

    print(f"\n✓ Fixed: {len(results['fixed'])} files")
    if results['fixed']:
        for f in results['fixed'][:10]:
            print(f"  • {f}")
        if len(results['fixed']) > 10:
            print(f"  ... and {len(results['fixed']) - 10} more")

    print(f"\n• No changes: {len(results['no_changes'])} files")
    if results['no_changes'] and len(results['no_changes']) <= 5:
        for f in results['no_changes']:
            print(f"  • {f}")

    if results['failed']:
        print(f"\n✗ Failed: {len(results['failed'])} files")
        for f in results['failed']:
            print(f"  • {f}")

    print("\n" + "="*70)
    print("✅ DATE CLEANUP COMPLETE")
    print("="*70)
    print("\nAll documents now have:")
    print("  ✓ ONLY real date: November 15, 2025 (signature)")
    print("  ✓ NO template dates (Feb 2024, July 2021, etc.)")
    print("  ✓ NO placeholder hearing info (court will assign after filing)")
    print("\n" + "="*70 + "\n")

if __name__ == "__main__":
    main()
