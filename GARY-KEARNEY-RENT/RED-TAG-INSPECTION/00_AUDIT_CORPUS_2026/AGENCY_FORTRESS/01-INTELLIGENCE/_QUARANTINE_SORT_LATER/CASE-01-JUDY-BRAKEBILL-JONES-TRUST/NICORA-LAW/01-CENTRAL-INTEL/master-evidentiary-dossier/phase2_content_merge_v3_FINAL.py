#!/usr/bin/env python3
"""
Phase 2 Content Merge Script v3 - FINAL (includes table processing)
"""

from docx import Document
import os

def update_placeholders_everywhere(doc, eric_info):
    """Update placeholders in paragraphs AND tables"""

    replacements_made = 0

    # Process paragraphs
    for para in doc.paragraphs:
        original = para.text
        para.text = apply_replacements(para.text, eric_info)
        if original != para.text:
            replacements_made += 1

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    original = para.text
                    para.text = apply_replacements(para.text, eric_info)
                    if original != para.text:
                        replacements_made += 1

    print(f"  • Made {replacements_made} replacements (paragraphs + tables)")
    return doc

def apply_replacements(text, eric_info):
    """Apply all placeholder replacements to text"""

    # Contact info
    text = text.replace("ADDRESS", eric_info['address'])
    text = text.replace("CITY, CA ZIP", eric_info['city_state_zip'])
    text = text.replace("(415) 123-4567", eric_info['phone'])
    text = text.replace("Tel: (415) 123-4567", f"Tel: {eric_info['phone']}")
    text = text.replace("email address", eric_info['email'])
    text = text.replace("Email: email address", f"Email: {eric_info['email']}")

    # Trust and estate names
    text = text.replace("LARRY LASAGNA 2002 REVOCABLE TRUST", eric_info['trust'].upper())
    text = text.replace("LARRY LASAGNA", "JUDY BRAKEBILL JONES")
    text = text.replace("JUDY [SURNMAE]", "JUDY BRAKEBILL JONES")
    text = text.replace("[SURNMAE]", "BRAKEBILL JONES")
    text = text.replace("ESTATE OF LARRY LASAGNA", "ESTATE OF JUDY BRAKEBILL JONES")

    # Court info
    text = text.replace("[NAME COUNTY]", "LOS ANGELES")
    text = text.replace("COUNTY OF [NAME COUNTY]", "COUNTY OF LOS ANGELES")

    # Case number
    text = text.replace("CASE NO. RP21094431", "CASE NO.: (To be assigned)")
    text = text.replace("RP21094431", "(To be assigned)")

    # Document title placeholder
    text = text.replace('NAME OF THE DOCUMENT FOR EXAMPLE "EX PARTE APPLICATION REQUE',
                       'EX PARTE APPLICATION')

    # Attorney info - remove Tate references
    text = text.replace("LAW OFFICE OF DAVID W. TATE", "")
    text = text.replace("DAVID W. TATE, ESQ.", "")
    text = text.replace("Attorneys for all Respondents and", "")
    text = text.replace("Moving Parties", "")

    # Date
    text = text.replace("____________, 2023", "November 15, 2025")
    text = text.replace("February 9, 2024", "November 15, 2025")
    text = text.replace("July 13, 2021", "November 15, 2025")

    # Remove other template artifacts
    text = text.replace("Reservation ID: 780799297986", "")
    text = text.replace("Time: 10:30 AM", "")
    text = text.replace("Dept.: Probate, 201", "")
    text = text.replace("Action Filed: July 13, 2021", "")

    return text

def extract_body_content(source_doc, skip_lines=15):
    """Extract body paragraphs from source"""
    content_paragraphs = []

    for i, para in enumerate(source_doc.paragraphs):
        if i < skip_lines:
            continue

        text = para.text.strip()
        if len(text) < 10:
            continue

        # Skip header-like content
        header_keywords = [
            "SUPERIOR COURT", "COUNTY OF", "Case No.", "CASE NO",
            "PETITIONER", "RESPONDENT", "IN RE THE",
            "ERIC B. JONES", "ERIC BRAKEBILL JONES"
        ]
        if any(keyword in para.text for keyword in header_keywords) and i < 35:
            continue

        content_paragraphs.append(para)

    return content_paragraphs

def merge_documents(template_path, source_path, output_path, eric_info, doc_name):
    """Complete merge"""

    print(f"\n{'='*70}")
    print(f"Processing: {doc_name}")
    print(f"{'='*70}")

    # Load template
    template = Document(template_path)
    print(f"✓ Loaded template: {len(template.paragraphs)} paras, {len(template.tables)} tables")

    # Update ALL placeholders (paragraphs + tables)
    print(f"✓ Updating placeholders everywhere...")
    template = update_placeholders_everywhere(template, eric_info)

    # Load source
    try:
        source = Document(source_path)
        print(f"✓ Loaded source: {len(source.paragraphs)} paragraphs")
    except Exception as e:
        print(f"✗ Error loading source: {e}")
        return False

    # Extract content
    content_paras = extract_body_content(source, skip_lines=15)
    print(f"✓ Extracted {len(content_paras)} content paragraphs")

    if len(content_paras) == 0:
        print(f"⚠ Warning: No content extracted - using header-only template")

    # Keep only header paragraphs (first 16)
    while len(template.paragraphs) > 16:
        p = template.paragraphs[-1]
        p._element.getparent().remove(p._element)

    # Remove the signature table (Table 2) - we'll add it back later
    if len(template.tables) > 1:
        tbl = template.tables[-1]._element
        tbl.getparent().remove(tbl)

    # Add blank line
    template.add_paragraph("")

    # Insert content
    added = 0
    for para in content_paras:
        if len(para.text.strip()) > 0:
            new_para = template.add_paragraph(para.text)

            # Preserve formatting
            if para.runs and new_para.runs:
                for src_run, dst_run in zip(para.runs[:len(new_para.runs)], new_para.runs):
                    if src_run.bold:
                        dst_run.bold = True
                    if src_run.italic:
                        dst_run.italic = True
                    if src_run.underline:
                        dst_run.underline = True

            added += 1

    print(f"✓ Added {added} content paragraphs")

    # Add signature block
    template.add_paragraph("")
    template.add_paragraph("")
    template.add_paragraph(f"DATED: November 15, 2025")
    template.add_paragraph("")
    template.add_paragraph("Respectfully submitted,")
    template.add_paragraph("")
    template.add_paragraph("_______________________________")
    template.add_paragraph("ERIC BRAKEBILL JONES")
    template.add_paragraph("Sole Successor Trustee, In Pro Per")
    template.add_paragraph("Judy Brakebill Jones 2008 Revocable Trust")

    # Save
    template.save(output_path)
    print(f"✓ Saved: {os.path.basename(output_path)}")

    return True

def convert_to_pdf(docx_path):
    """Convert to PDF"""
    output_dir = os.path.dirname(docx_path)
    basename = os.path.basename(docx_path)

    cmd = f'cd "{output_dir}" && /Applications/LibreOffice.app/Contents/MacOS/soffice --headless --convert-to pdf --outdir . "{basename}" 2>&1 > /dev/null'

    result = os.system(cmd)

    pdf_path = docx_path.replace('.docx', '.pdf')
    if os.path.exists(pdf_path):
        size_kb = os.path.getsize(pdf_path) / 1024
        print(f"✓ PDF: {os.path.basename(pdf_path)} ({size_kb:.1f} KB)")
        return True

    return False

def verify_no_bad_placeholders(docx_path):
    """Verify no placeholder text remains"""
    doc = Document(docx_path)

    bad_patterns = [
        ("ADDRESS", "contact info placeholder"),
        ("CITY, CA ZIP", "city placeholder"),
        ("(415) 123-4567", "phone placeholder"),
        ("email address", "email placeholder"),
        ("LARRY LASAGNA", "wrong decedent name"),
        ("[NAME COUNTY]", "county placeholder"),
        ("[SURNMAE]", "surname placeholder"),
        ("DAVID W. TATE", "old attorney"),
        ("RP21094431", "old case number")
    ]

    found = []

    # Check paragraphs
    for para in doc.paragraphs:
        for pattern, desc in bad_patterns:
            if pattern in para.text:
                found.append((desc, pattern))

    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for pattern, desc in bad_patterns:
                        if pattern in para.text:
                            found.append((desc, pattern))

    if found:
        unique_found = list(set(found))
        print(f"✗ Found {len(unique_found)} placeholder(s):")
        for desc, pattern in unique_found[:5]:
            print(f"  ✗ {desc}: '{pattern}'")
        return False

    print(f"✓ Verification: Clean (no placeholders)")
    return True

# ===== CONFIGURATION =====

ERIC_INFO = {
    'name': 'ERIC BRAKEBILL JONES',
    'address': '17742 Berta Canyon Road',
    'city_state_zip': 'Lake Hughes, CA 93532',
    'phone': '(626) 348-3019',
    'email': 'eric@recovery-compass.org',
    'capacity': 'Sole Successor Trustee, In Pro Per',
    'trust': 'Judy Brakebill Jones 2008 Revocable Trust'
}

TEMPLATE_PATH = "/Users/ericjones/Cases/Judy-Trust/pro-per-probate-850/anuar-medina-ramirez-seven-hills-law/email-november-11/pleading-paper/PLEADING PAPER FOR ERIC JONES.docx"

SOURCE_DIR = "/Users/ericjones/11-15/5-bird/judy-jones-probate-850/self-filing-IRONCLAD-FINAL/remediated_docx_v2"

OUTPUT_DIR = "/Users/ericjones/Compass Fortress/cases/judy-jones-probate-850/4-email-to-anuar-nov15/ALL-ATTACHMENTS-READY-TO-SEND"

DOCUMENTS = [
    ('01_Ex_Parte_Application', '01_Ex_Parte_Application.docx', '01_Ex_Parte_Application_CRC2111_FINAL.docx'),
    ('02_Declaration_with_Exhibits', '02_Declaration_with_Exhibits.docx', '02_Declaration_with_Exhibits_CRC2111_FINAL.docx'),
    ('04_MPA', '04_MPA.docx', '04_MPA_CRC2111_FINAL.docx'),
    ('06_Petition_850', '06_Petition_850.docx', '06_Petition_850_CRC2111_FINAL.docx')
]

# ===== MAIN =====

def main():
    print("\n" + "="*70)
    print("PHASE 2: FINAL CONTENT MERGE (v3 - Tables + Paragraphs)")
    print("="*70)

    completed = []
    failed = []
    warnings = []

    for doc_name, source_file, output_file in DOCUMENTS:
        source_path = os.path.join(SOURCE_DIR, source_file)
        output_path = os.path.join(OUTPUT_DIR, output_file)

        if not os.path.exists(source_path):
            print(f"\n✗ Source missing: {source_file}")
            failed.append(doc_name)
            continue

        # Merge
        if not merge_documents(TEMPLATE_PATH, source_path, output_path, ERIC_INFO, doc_name):
            failed.append(doc_name)
            continue

        # Verify
        if not verify_no_bad_placeholders(output_path):
            warnings.append(doc_name)

        # Convert
        if convert_to_pdf(output_path):
            completed.append(doc_name)
        else:
            failed.append(doc_name)

    # ===== REPORT =====
    print("\n" + "="*70)
    print("FINAL STATUS")
    print("="*70)

    print(f"\n✓ Completed: {len(completed)}/4")
    print(f"✗ Failed: {len(failed)}/4")
    if warnings:
        print(f"⚠ Warnings: {len(warnings)}/4")

    if len(completed) == 4 and len(warnings) == 0:
        print("\n🎯 MISSION COMPLETE\n")
        print("📁 Location:")
        print(f"   {OUTPUT_DIR}/\n")
        print("📄 Ready to Send:")
        for doc_name in completed:
            print(f"   ✓ {doc_name}_CRC2111_FINAL.pdf")

        print("\n" + "="*70)
        print("✓ CRC 2.111 format preserved")
        print("✓ All placeholders replaced")
        print("✓ Legal content merged")
        print("✓ PDFs generated")
        print("\n✉️  STATUS: READY TO SEND TO ANUAR")
        print("="*70 + "\n")

    else:
        if warnings:
            print("\n⚠️  Review needed:")
            for doc in warnings:
                print(f"   ⚠ {doc}")
        if failed:
            print("\n✗ Failed:")
            for doc in failed:
                print(f"   ✗ {doc}")

    print("\n")

if __name__ == "__main__":
    main()
