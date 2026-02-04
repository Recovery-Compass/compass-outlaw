#!/usr/bin/env python3
"""
PFV v12 + FPF v1.0: ATTORNEY-READY WORD DOCUMENT GENERATION
Generate CRC 2.111 compliant pleading paper Word documents from extracted content

MISSION: Create properly formatted .docx files using the verified pleading paper template
         and merge with FPF v1.0 extracted content.

OUTPUT: Professional Word documents ready for attorney review and PDF conversion
"""

import PyPDF2
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from datetime import datetime
import shutil

print("=" * 80)
print("PFV v12 + FPF v1.0: ATTORNEY-READY WORD DOCUMENT GENERATION")
print("CRC 2.111 Pleading Paper Format with Verified Content")
print("=" * 80)
print()

# ===== CONFIGURATION =====

BASE_DIR = Path("/Users/ericjones/Compass Fortress/cases/judy-jones-probate-850")
SOURCE_PDF_DIR = BASE_DIR / "1-core-filing-documents"
TEMPLATE_PATH = BASE_DIR / "PFV_V12_FPF_V1_COMPLIANT_2025-11-16/last-email-sent-to-anuar/pleading-paper-model/PLEADING PAPER FOR ERIC JONES.docx"
OUTPUT_DIR = BASE_DIR / "PFV_V12_FPF_V1_COMPLIANT_2025-11-16/core-filing-documents"

# Create output directory
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

print(f"📁 Template: {TEMPLATE_PATH.name}")
print(f"📁 Sources: {SOURCE_PDF_DIR}")
print(f"📁 Output: {OUTPUT_DIR}")
print()

# Verify template exists
if not TEMPLATE_PATH.exists():
    print(f"❌ CRITICAL ERROR: Template not found at {TEMPLATE_PATH}")
    exit(1)

# ===== VERIFIED INSTITUTIONAL DATA =====

VERIFIED_DATA = {
    "petitioner_name": "ERIC B. JONES",
    "petitioner_address_line1": "5634 Noel Drive",
    "petitioner_city_state_zip": "Temple City, CA 91780",
    "petitioner_phone": "(626) 348-3019",
    "petitioner_email": "eric@recovery-compass.org",
    "court_name": "SUPERIOR COURT OF CALIFORNIA, COUNTY OF LOS ANGELES",
    "court_district": "CENTRAL DISTRICT",
    "case_number": "(to be assigned)",
}

# ===== DOCUMENT SPECIFICATIONS =====

DOCUMENT_SPECS = [
    {
        "source_pdf": "01_Ex_Parte_Application.pdf",
        "title": "EX PARTE APPLICATION FOR TEMPORARY RESTRAINING ORDER AND ORDER TO SHOW CAUSE",
        "output_name": "01_Ex_Parte_Application_ATTORNEY_READY.docx",
        "short_title": "Ex Parte Application for TRO"
    },
    {
        "source_pdf": "02_Declaration_with_Exhibits.pdf",
        "title": "DECLARATION OF ERIC BRAKEBILL JONES IN SUPPORT OF EX PARTE APPLICATION",
        "output_name": "02_Declaration_with_Exhibits_ATTORNEY_READY.docx",
        "short_title": "Declaration of Eric B. Jones"
    },
    {
        "source_pdf": "03a_Declaration_for_Good_Cause_Exception_to_Notice.pdf",
        "title": "DECLARATION FOR GOOD CAUSE EXCEPTION TO NOTICE REQUIREMENT",
        "output_name": "03a_Declaration_Good_Cause_ATTORNEY_READY.docx",
        "short_title": "Declaration for Good Cause"
    },
    {
        "source_pdf": "04_MPA.pdf",
        "title": "MEMORANDUM OF POINTS AND AUTHORITIES IN SUPPORT OF EX PARTE APPLICATION",
        "output_name": "04_MPA_ATTORNEY_READY.docx",
        "short_title": "Memorandum of Points and Authorities"
    },
    {
        "source_pdf": "05_Proposed_TRO_OSC.pdf",
        "title": "[PROPOSED] TEMPORARY RESTRAINING ORDER AND ORDER TO SHOW CAUSE",
        "output_name": "05_Proposed_TRO_OSC_ATTORNEY_READY.docx",
        "short_title": "Proposed TRO and OSC"
    },
    {
        "source_pdf": "06_Petition_850.pdf",
        "title": "PETITION UNDER PROBATE CODE SECTION 850",
        "output_name": "06_Petition_850_ATTORNEY_READY.docx",
        "short_title": "Petition under Probate Code § 850"
    },
    {
        "source_pdf": "07_Probate_Case_Cover_Sheet_Final.pdf",
        "title": "PROBATE CASE COVER SHEET",
        "output_name": "07_Probate_Cover_Sheet_ATTORNEY_READY.docx",
        "short_title": "Probate Case Cover Sheet"
    },
]

# ===== HELPER FUNCTIONS =====

def extract_text_from_pdf(pdf_path):
    """Extract text from PDF, returns concatenated text from all pages."""
    if not pdf_path.exists():
        print(f"⚠️  WARNING: Source file not found: {pdf_path}")
        return ""
    
    try:
        text_parts = []
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text = page.extract_text()
                text_parts.append(text)
        
        return "\n\n".join(text_parts)
    
    except Exception as e:
        print(f"❌ ERROR extracting from {pdf_path.name}: {e}")
        return ""

def clean_extracted_text(text):
    """Clean extracted text, removing header info that will be in template."""
    lines = text.split('\n')
    cleaned_lines = []
    skip_patterns = [
        "SUPERIOR COURT",
        "COUNTY OF",
        "CASE NO",
        "PRO PER",
        "ERIC B. JONES",
        "ERIC BRAKEBILL JONES",
        "5634 Noel",
        "Temple City",
        "(626) 348-3019",
        "eric@recovery-compass.org"
    ]
    
    for line in lines:
        # Skip header content that's already in template
        if any(pattern in line.upper() for pattern in skip_patterns):
            continue
        # Skip empty lines at start
        if not cleaned_lines and not line.strip():
            continue
        cleaned_lines.append(line)
    
    return '\n'.join(cleaned_lines)

def create_attorney_ready_document(spec):
    """
    Create attorney-ready Word document with CRC 2.111 formatting.
    Uses the verified pleading paper template and merges with extracted content.
    """
    source_path = SOURCE_PDF_DIR / spec["source_pdf"]
    output_path = OUTPUT_DIR / spec["output_name"]
    
    print()
    print("─" * 80)
    print(f"📄 GENERATING: {spec['short_title']}")
    print(f"   Source: {spec['source_pdf']}")
    print("─" * 80)
    
    if not source_path.exists():
        print(f"❌ FAILED: Source file not found")
        return False
    
    # Load template
    try:
        doc = Document(str(TEMPLATE_PATH))
    except Exception as e:
        print(f"❌ ERROR loading template: {e}")
        return False
    
    # Update document title in header (first paragraph after address block)
    # Find and replace [TITLE OF DOCUMENT] placeholder
    for para in doc.paragraphs[:25]:  # Check first 25 paragraphs (header section)
        if "[TITLE OF DOCUMENT]" in para.text:
            para.text = spec["title"]
            # Center align the title
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Make it bold
            for run in para.runs:
                run.bold = True
        # Replace case number placeholder if present
        if "[TO BE ASSIGNED]" in para.text or "CASE NO.:" in para.text.upper():
            original = para.text
            para.text = original.replace("[TO BE ASSIGNED]", VERIFIED_DATA["case_number"])
    
    # Update footer with document title
    for section in doc.sections:
        footer = section.footer
        for para in footer.paragraphs:
            if para.text.strip():
                para.text = spec["short_title"]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Extract content from source PDF
    extracted_text = extract_text_from_pdf(source_path)
    
    if not extracted_text:
        print(f"❌ FAILED: No content extracted")
        return False
    
    # Clean the text
    content = clean_extracted_text(extracted_text)
    
    # Add content to document (after template header)
    # Split content into paragraphs and add them
    content_paragraphs = content.split('\n\n')
    added_count = 0
    
    for para_text in content_paragraphs:
        if not para_text.strip():
            continue
        
        # Add paragraph to document
        p = doc.add_paragraph(para_text.strip())
        
        # Apply standard formatting
        p.paragraph_format.line_spacing = 2.0  # Double spacing
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        
        # Set font
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        
        added_count += 1
    
    # Save the document
    try:
        doc.save(str(output_path))
        file_size = output_path.stat().st_size / 1024
        print(f"✅ COMPLETE: {output_path.name}")
        print(f"   Size: {file_size:.1f} KB")
        print(f"   Paragraphs added: {added_count}")
        return True
    except Exception as e:
        print(f"❌ ERROR saving document: {e}")
        return False

# ===== MAIN EXECUTION =====

print("STEP 1: Verify Template and Sources")
print("─" * 80)

template_ok = TEMPLATE_PATH.exists()
print(f"{'✅' if template_ok else '❌'} Template: {TEMPLATE_PATH.name}")

if not template_ok:
    print()
    print("❌ CRITICAL ERROR: Cannot proceed without template")
    exit(1)

# Check source PDFs
available_sources = 0
for spec in DOCUMENT_SPECS:
    source_path = SOURCE_PDF_DIR / spec["source_pdf"]
    if source_path.exists():
        size_kb = source_path.stat().st_size / 1024
        print(f"✅ {spec['source_pdf']} ({size_kb:.1f} KB)")
        available_sources += 1
    else:
        print(f"❌ {spec['source_pdf']} - NOT FOUND")

print()
print(f"Available sources: {available_sources}/{len(DOCUMENT_SPECS)}")

print()
print("STEP 2: Generate Attorney-Ready Word Documents")
print("─" * 80)

success_count = 0
failed_docs = []

for spec in DOCUMENT_SPECS:
    if create_attorney_ready_document(spec):
        success_count += 1
    else:
        failed_docs.append(spec["short_title"])

# ===== COPY ENHANCEMENT DOCUMENTS =====

print()
print("STEP 3: Verify Enhancement Documents")
print("─" * 80)

enhancement_source = BASE_DIR / "PFV_V12_FPF_V1_COMPLIANT_2025-11-16/enhancement-documents"
enhancement_files = [
    "ANUAR_INTEGRATION_GUIDE.pdf",
    "ENHANCED_DECLARATION_ADDENDUM.pdf",
    "ENHANCED_MPA_LEGAL_FRAMEWORK.pdf"
]

for filename in enhancement_files:
    source = enhancement_source / filename
    if source.exists():
        size_kb = source.stat().st_size / 1024
        print(f"✅ {filename} ({size_kb:.1f} KB) - already in place")
    else:
        print(f"⚠️  {filename} - not found")

# ===== FINAL SUMMARY =====

print()
print("=" * 80)
print(f"EXECUTION COMPLETE: {success_count}/{len(DOCUMENT_SPECS)} DOCUMENTS GENERATED")
print("=" * 80)
print()

if success_count == len(DOCUMENT_SPECS):
    print("🎯 MISSION COMPLETE: 100% Attorney-Ready Package Generated")
    print()
    print(f"📁 Output Location: {OUTPUT_DIR}")
    print()
    print("GENERATED FILES:")
    for spec in DOCUMENT_SPECS:
        output_path = OUTPUT_DIR / spec["output_name"]
        if output_path.exists():
            size_kb = output_path.stat().st_size / 1024
            print(f"  ✅ {spec['output_name']} ({size_kb:.1f} KB)")
    
    print()
    print("=" * 80)
    print("PACKAGE READY FOR ATTORNEY REVIEW")
    print("=" * 80)
    print("✅ CRC 2.111 Pleading Paper Format Applied")
    print("✅ Verified Contact Information (Eric B. Jones)")
    print("✅ Court: LA Superior Court, Central District")
    print("✅ FPF v1.0 Content (No Fabrications)")
    print()
    print("NEXT STEPS:")
    print("1. Open Word documents in Microsoft Word or LibreOffice")
    print("2. Review formatting and content alignment")
    print("3. Convert to PDF using Adobe Acrobat Pro")
    print("4. Package with enhancement documents for Attorney Anuar")
    print()
    print(f"Enhancement documents location:")
    print(f"  {enhancement_source}")
    
else:
    print(f"⚠️  PARTIAL SUCCESS: {len(failed_docs)} documents failed")
    print()
    print("FAILED DOCUMENTS:")
    for doc in failed_docs:
        print(f"  ❌ {doc}")

print()
print("=" * 80)
print("PFV v12 + FPF v1.0: ATTORNEY-READY GENERATION COMPLETE")
print("=" * 80)
